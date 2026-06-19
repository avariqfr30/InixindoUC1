"""Regression coverage for DOCX rendering hygiene."""
from __future__ import annotations

import sys
import unittest
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))


class DocumentRenderingHygieneTest(unittest.TestCase):
    def test_renderer_strips_escaped_source_markup_and_demotes_invalid_pipe_rows(self) -> None:
        from main.document_rendering import DocumentBuilder

        doc = Document()
        raw = (
            "AccelByte membutuhkan tata kelola layanan digital yang selaras dengan operasi backend game "
            "&lt;/em&gt;&lt;/span&gt;.\n\n"
            "| Sumber eksternal 1 | fakta=raw snippet | sumber=Example | url=https://example.test |\n\n"
            "Narasi dilanjutkan sebagai paragraf biasa."
        )

        DocumentBuilder.process_content(doc, raw, (0, 51, 102), "BAB I")
        rendered_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)

        self.assertIn("AccelByte membutuhkan tata kelola layanan digital", rendered_text)
        self.assertIn("Narasi dilanjutkan", rendered_text)
        self.assertNotIn("span", rendered_text.lower())
        self.assertNotIn("fakta=", rendered_text)
        self.assertEqual(len(doc.tables), 0)

    def test_renderer_keeps_valid_markdown_tables(self) -> None:
        from main.document_rendering import DocumentBuilder

        doc = Document()
        raw = (
            "| Area | Relevansi |\n"
            "| --- | --- |\n"
            "| Governance | Mengawal keputusan dan risiko |\n"
        )

        DocumentBuilder.process_content(doc, raw, (0, 51, 102), "BAB XI")

        self.assertEqual(len(doc.tables), 1)
        self.assertEqual(doc.tables[0].cell(1, 0).text.strip(), "Governance")

    def test_renderer_deduplicates_repeated_markdown_table_body_rows(self) -> None:
        from main.document_rendering import DocumentBuilder

        doc = Document()
        raw = (
            "| Peran | Fokus | Dampak |\n"
            "| --- | --- | --- |\n"
            "| Project Manager | Mengawal scope | Keputusan terkendali |\n"
            "| Project Manager | Mengawal scope | Keputusan terkendali |\n"
            "| Tenaga Ahli | Menyusun arsitektur | Desain lebih siap |\n"
        )

        DocumentBuilder.process_content(doc, raw, (0, 51, 102), "BAB XI")

        self.assertEqual(len(doc.tables), 1)
        self.assertEqual(len(doc.tables[0].rows), 3)

    def test_bullets_after_markdown_table_do_not_become_table_rows(self) -> None:
        from main.document_rendering import DocumentBuilder

        doc = Document()
        raw = (
            "| Peran Tenaga Ahli | Fokus Tanggung Jawab | Kompetensi / Pengalaman Kunci | Keterlibatan |\n"
            "| --- | --- | --- | --- |\n"
            "| Project Manager | Mengawal scope | Governance dan quality gate | Aktif |\n"
            "- Untuk Ajinomoto Indonesia, pembahasan tetap diarahkan pada keputusan dan langkah kerja yang konkret.\n"
            "- KPI acuan yang tetap dijaga adalah outcome bisnis utama bergerak positif.\n"
        )

        DocumentBuilder.process_content(doc, raw, (0, 51, 102), "BAB XI")
        rendered_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)

        self.assertEqual(len(doc.tables), 1)
        self.assertEqual(len(doc.tables[0].rows), 2)
        self.assertIn("Untuk Ajinomoto Indonesia", rendered_text)
        self.assertIn("KPI acuan", rendered_text)

    def test_flattened_markdown_separator_tokens_do_not_render_as_prose(self) -> None:
        from main.document_rendering import DocumentBuilder

        doc = Document()
        raw = (
            "Masalah utama: Risiko Utama Dampak Bisnis Dampak Operasional Respon yang Diperlukan "
            "--- --- --- --- Prioritas berubah-ubah membuat sponsor sulit membaca keputusan. "
            "Kapabilitas pelaksana Nama Posisi Utama Kapabilitas ---."
        )

        DocumentBuilder.process_content(doc, raw, (0, 51, 102), "BAB II")
        rendered_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)

        self.assertIn("Masalah utama", rendered_text)
        self.assertIn("Prioritas berubah-ubah", rendered_text)
        self.assertNotIn("---", rendered_text)

    def test_writer_profile_summarizes_public_approach_and_certifications(self) -> None:
        from main.document_rendering import DocumentBuilder

        doc = Document()
        profile = {
            "profile_summary": "Ringkasan profil.",
            "credential_highlights": (
                "Kapabilitas internal didukung tenaga ahli bernama, termasuk Andi Yuniantoro "
                "(Project Manager dengan sertifikasi CAPM, TOGAF 9 Foundations, COBIT 5, ITIL Foundation V3), "
                "Citra Arfanudin (Tenaga Ahli dengan sertifikasi Lead Auditor ISO 27001, CEH)."
            ),
            "values_approach": " ".join(["Pendekatan publik terlalu panjang"] * 30),
            "certifications": "ISO, Lihat website resmi untuk daftar lengkap",
        }

        DocumentBuilder.add_writer_firm_profile_section(doc, profile, (0, 51, 102))
        text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
        table_text = "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)

        self.assertNotIn("Lihat website resmi", table_text)
        self.assertIn("CAPM", table_text)
        self.assertIn("TOGAF", table_text)
        self.assertNotIn("Andi Yuniantoro", table_text)
        self.assertIn("Ringkasan profil", text)


if __name__ == "__main__":
    unittest.main()
