import sys
import unittest
import zipfile
from io import BytesIO
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from main.document_rendering import DocumentBuilder
from main.executive_summary import ExecutiveSummaryBuilder
from main.finance import FinancialAnalyzer
from main.reader_facing_hygiene import sanitize_reader_facing_sources
from main.config import UNIVERSAL_STRUCTURE


class ReaderFacingDocumentContractTests(unittest.TestCase):
    def test_standard_proposal_chapters_keep_scope_and_post_flowchart_closing(self):
        expected = [
            ("c_1", "BAB I – KONTEKS KLIEN"),
            ("c_2", "BAB II – PERMASALAHAN"),
            ("c_7", "BAB III – RUANG LINGKUP PEKERJAAN"),
            ("c_3", "BAB IV – KLASIFIKASI KEBUTUHAN"),
            ("c_4", "BAB V – PENDEKATAN"),
            ("c_5", "BAB VI – METODOLOGI"),
            ("c_6", "BAB VII – SOLUTION DESIGN"),
            ("c_8", "BAB VIII – TIMELINE PEKERJAAN"),
            ("c_9", "BAB IX – TATA KELOLA PROYEK"),
            ("c_11", "BAB X – STRUKTUR & TEAM PROYEK"),
            ("c_12", "BAB XI – MODEL PEMBIAYAAN"),
            ("c_closing", "BAB XII – PENUTUP"),
        ]

        self.assertEqual([(chapter["id"], chapter["title"]) for chapter in UNIVERSAL_STRUCTURE], expected)
        scope_chapter = UNIVERSAL_STRUCTURE[2]
        self.assertIn("membatasi pekerjaan", scope_chapter["length_intent"].lower())
        self.assertIn("off-scope", scope_chapter["length_intent"].lower())
        closing = UNIVERSAL_STRUCTURE[-1]
        self.assertIn("post-flowchart", closing["keywords"])
        self.assertIn("email", closing["length_intent"].lower())

    def test_team_chapter_is_the_writer_firm_capability_home(self):
        team_chapter = next(chapter for chapter in UNIVERSAL_STRUCTURE if chapter["id"] == "c_11")

        self.assertEqual(team_chapter["title"], "BAB X – STRUKTUR & TEAM PROYEK")
        self.assertIn("kapabilitas konsultan", team_chapter["length_intent"].lower())
        self.assertIn("sertifikasi", team_chapter["length_intent"].lower())
        self.assertIn("profil perusahaan penyusun", team_chapter["length_intent"].lower())

    def test_sanitizer_removes_internal_source_mode_terms(self):
        raw = (
            "Proposal memakai APIDog, Internal API endpoint /api/Resource/dataset, "
            "dataset code ReferenceDataset dan dataset name ConsultantProjectExpertHistory, "
            "source-of-truth cache sync, agent workflow, evidence card, confidence label, "
            "[RESEARCH_AGENT] [INTERNAL_DATA_AGENT] [COMMERCIAL_STRATEGY_AGENT] "
            "[TECHNICAL_SOLUTION_AGENT] [RISK_COMPLIANCE_AGENT] [EDITOR_MAIN_AGENT], "
            "Data Internal ReferenceAccount, dan URL https://example.test/api."
        )
        clean = sanitize_reader_facing_sources(raw)
        forbidden_tokens = [
            "APIDog",
            "Internal API",
            "endpoint",
            "/api/Resource/dataset",
            "dataset",
            "dataset code",
            "dataset name",
            "ReferenceDataset",
            "ReferenceAccount",
            "ConsultantProjectExpertHistory",
            "source-of-truth",
            "cache",
            "sync",
            "agent",
            "workflow",
            "evidence card",
            "confidence label",
            "RESEARCH_AGENT",
            "INTERNAL_DATA_AGENT",
            "COMMERCIAL_STRATEGY_AGENT",
            "TECHNICAL_SOLUTION_AGENT",
            "RISK_COMPLIANCE_AGENT",
            "EDITOR_MAIN_AGENT",
            "Data Internal",
            "https://",
        ]
        for token in forbidden_tokens:
            self.assertNotIn(token, clean)
        self.assertIn("konteks yang tersedia", clean)

    def test_synthesizer_blocks_raw_ui_and_api_helper_inputs(self):
        raw = (
            "Nama Perusahaan Klien: ReferenceAccount mencatat source=/api/Resource/dataset "
            "dan dataset_code=ConsultantProjectExpertHistory. Dirangkum dari sumber APIDog: "
            "Problem, Opportunity, Directive untuk Pain Points."
        )

        clean = sanitize_reader_facing_sources(raw)

        forbidden_tokens = [
            "Nama Perusahaan Klien",
            "ReferenceAccount",
            "source=",
            "/api/Resource/dataset",
            "dataset_code",
            "ConsultantProjectExpertHistory",
            "Dirangkum dari sumber",
            "APIDog",
            "Problem, Opportunity, Directive",
            "Pain Points",
        ]
        for token in forbidden_tokens:
            self.assertNotIn(token, clean)
        self.assertIn("catatan klien", clean)
        self.assertIn("riwayat pengalaman konsultan", clean)
        self.assertIn("kebutuhan prioritas yang perlu dipertegas", clean)

    def test_toc_does_not_emit_static_fallback_items(self):
        doc = Document()
        DocumentBuilder.add_table_of_contents(doc, ["Ringkasan Eksekutif", "BAB I - Pendekatan", "BAB II - Rencana Kerja"])
        text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
        self.assertIn("DAFTAR ISI", text)
        self.assertNotIn("Ringkasan Eksekutif", text)
        self.assertNotIn("BAB I - Pendekatan", text)
        self.assertNotIn("BAB II - Rencana Kerja", text)
        self.assertNotIn("Update Field", text)

    def test_toc_uses_word_field_and_updates_on_open(self):
        doc = Document()
        DocumentBuilder.add_table_of_contents(doc, ["Ringkasan Eksekutif", "BAB I - Pendekatan"])

        output = BytesIO()
        doc.save(output)
        output.seek(0)

        with zipfile.ZipFile(output) as archive:
            document_xml = archive.read("word/document.xml").decode("utf-8")
            settings_xml = archive.read("word/settings.xml").decode("utf-8")

        self.assertIn('TOC \\o "1-3" \\h \\z \\u', document_xml)
        self.assertIn('w:fldCharType="begin"', document_xml)
        self.assertIn('w:fldCharType="separate"', document_xml)
        self.assertIn('w:fldCharType="end"', document_xml)
        self.assertIn("w:updateFields", settings_xml)
        self.assertIn('w:val="true"', settings_xml)

    def test_sanitizer_translates_unnecessary_english_reader_labels(self):
        raw = (
            "Executive Summary dan Key Findings menampilkan Recommendation, "
            "Dashboard, Insight, Pain Points, Current State, Target State, "
            "Owner, Scope, Deliverables, Milestones, Roadmap, dan Quality Gate. "
            "Brand teknis seperti API, OSINT, RAG, UAT, Go-Live, Microsoft Azure, dan Power BI tetap relevan."
        )
        clean = sanitize_reader_facing_sources(raw)
        forbidden_tokens = [
            "Executive Summary",
            "Key Findings",
            "Recommendation",
            "Dashboard",
            "Insight",
            "Pain Points",
            "Current State",
            "Target State",
            "Owner",
            "Scope",
            "Deliverables",
            "Milestones",
            "Roadmap",
            "Quality Gate",
        ]
        for token in forbidden_tokens:
            self.assertNotIn(token, clean)
        for expected in [
            "Ringkasan Eksekutif",
            "Temuan Utama",
            "Rekomendasi",
            "dasbor",
            "wawasan",
            "Titik Masalah",
            "kondisi saat ini",
            "kondisi target",
            "penanggung jawab",
            "ruang lingkup",
            "keluaran kerja",
            "tonggak kerja",
            "peta jalan",
            "gerbang mutu",
        ]:
            self.assertIn(expected, clean)
        for expected in ["*API*", "*OSINT*", "*RAG*", "*UAT*", "*Go-Live*", "Microsoft Azure", "Power BI"]:
            self.assertIn(expected, clean)

    def test_executive_summary_is_decision_first_and_source_neutral(self):
        summary = ExecutiveSummaryBuilder.build(
            client="PT Contoh",
            project="Implementasi AI Adoption",
            project_goal="meningkatkan produktivitas layanan",
            timeline="12 minggu",
            budget="Rp500 juta",
            value_map={
                "primary_outcome": "adopsi berjalan terukur",
                "value_statement": "produktivitas layanan meningkat tanpa mengorbankan tata kelola",
                "proof_points": ["pengalaman implementasi tata kelola digital", "indikator keberhasilan disepakati sejak awal"],
            },
            personalization_pack={
                "industry": "Teknologi Layanan Digital",
                "relationship_mode": "new",
                "kpi_blueprint": ["waktu respons layanan", "adopsi proses kerja"],
            },
        )
        self.assertIn("# Ringkasan Eksekutif", summary)
        expected_sections = [
            "## Inti Keputusan",
            "## Situasi dan Masalah Klien",
            "## Solusi yang Direkomendasikan",
            "## Nilai dan Bukti",
            "## Prioritas Eksekusi",
            "## Risiko yang Perlu Dikendalikan",
            "## Keputusan Berikutnya",
        ]
        positions = [summary.index(section) for section in expected_sections]
        self.assertEqual(positions, sorted(positions))
        self.assertIn("PT Contoh", summary)
        self.assertIn("Implementasi AI Adoption", summary)
        self.assertIn("Teknologi Layanan Digital", summary)
        self.assertIn("12 minggu", summary)
        self.assertIn("Rp500 juta", summary)
        for token in [
            "APIDog",
            "Internal API",
            "Data Internal",
            "dataset",
            "source-of-truth",
            "cache",
            "sync",
            "agent",
            "workflow",
            "evidence card",
            "confidence label",
            "method",
            "source",
            "process",
            "proses internal",
        ]:
            self.assertNotIn(token, summary)
        for heading in [
            "## BLUF",
            "## Bottom Line",
            "## Key Findings",
            "## Recommendation",
            "## Executive Summary",
            "## Dashboard",
            "## Insight",
        ]:
            self.assertNotIn(heading, summary)

    def test_executive_summary_is_synthesized_from_finished_chapters(self):
        chapters = [
            {"id": "c_7", "title": "BAB III – RUANG LINGKUP PEKERJAAN"},
            {"id": "c_6", "title": "BAB VII – SOLUTION DESIGN"},
            {"id": "c_12", "title": "BAB XI – MODEL PEMBIAYAAN"},
        ]
        chapter_outputs = {
            "c_7": (
                "## 3.1 Lingkup Pekerjaan Utama\n"
                "Pekerjaan mencakup asesmen tata kelola layanan digital dan penyusunan roadmap prioritas.\n\n"
                "## 3.2 Batasan Pekerjaan, Asumsi, dan Hal di Luar Cakupan\n"
                "Implementasi penuh aplikasi, pengadaan lisensi, dan pekerjaan di luar change request tertulis berada di luar cakupan."
            ),
            "c_6": "## 7.1 Solusi/Output Metodologi yang Dibangun\nKeluaran utama adalah dokumen rekomendasi, matriks prioritas, dan rencana implementasi bertahap.",
            "c_12": "## 11.1 Biaya dan Tahapan Pembayaran\nModel pembiayaan mengikuti milestone dan finalisasi scope.",
        }

        summary = ExecutiveSummaryBuilder.build_from_chapters(
            client="PT Contoh",
            project="Transformasi Layanan Digital",
            project_goal="tata kelola layanan digital",
            timeline="12 minggu",
            budget="Rp500 juta",
            value_map={},
            personalization_pack={},
            selected_chapters=chapters,
            chapter_outputs=chapter_outputs,
        )

        self.assertIn("# Ringkasan Eksekutif", summary)
        self.assertIn("asesmen tata kelola layanan digital", summary)
        self.assertIn("Implementasi penuh aplikasi", summary)
        self.assertIn("dokumen rekomendasi", summary)
        self.assertIn("tonggak kerja", summary)
        self.assertLess(len(summary.split()), 520)

    def test_budget_recommendation_uses_silent_scope_contract(self):
        narrow = FinancialAnalyzer._dynamic_budget_from_osint(
            "PT Contoh",
            finance_snippets=[],
            benchmark_snippets=[],
            timeline="3 bulan",
            project_type="Strategic",
            service_type="Konsultan",
            project_goal="Problem",
            objective="menyusun roadmap tata kelola layanan digital",
            notes="asesmen dan rekomendasi prioritas",
            frameworks="ISO",
            scope_context="In-scope: asesmen, workshop, roadmap. Out-of-scope: implementasi penuh aplikasi, lisensi, integrasi sistem.",
        )
        broad = FinancialAnalyzer._dynamic_budget_from_osint(
            "PT Contoh",
            finance_snippets=[],
            benchmark_snippets=[],
            timeline="3 bulan",
            project_type="Strategic",
            service_type="Konsultan",
            project_goal="Problem",
            objective="menyusun roadmap tata kelola layanan digital",
            notes="asesmen, implementasi penuh aplikasi, integrasi sistem, migrasi data, dan pendampingan go-live",
            frameworks="ISO, COBIT, Regulasi",
            scope_context="In-scope: asesmen, implementasi penuh aplikasi, integrasi sistem, migrasi data, pelatihan, dan go-live.",
        )

        def standard_price(payload):
            raw = next(item["price"] for item in payload["options"] if item["tier"] == "Standard")
            return int("".join(ch for ch in raw if ch.isdigit()))

        self.assertIn("ruang lingkup", narrow["analysis"].lower())
        self.assertGreater(standard_price(broad), standard_price(narrow))

    def test_executive_summary_does_not_echo_internal_helper_or_bab_titles(self):
        summary = ExecutiveSummaryBuilder.build(
            client="Accelbyte",
            project=(
                "Konteks akun internal menempatkan Accelbyte dengan berlokasi di Daerah Istimewa Yogyakarta; "
                "segmen swasta / swasta. Gunakan informasi ini sebagai latar segmentasi dan lokasi, bukan sebagai rumusan tujuan proyek."
            ),
            project_goal="Problem, Opportunity, Directive",
            timeline="4 Bulan",
            budget="Rp 192.222.720",
            value_map={
                "primary_outcome": "tata kelola layanan digital yang lebih terarah",
                "value_statement": (
                    "membantu sponsor mengubah kebutuhan awal menjadi prioritas, keluaran kerja, dan ukuran keberhasilan"
                ),
            },
            personalization_pack={
                "industry": "Game Technology & Digital Platform",
                "kpi_blueprint": ["kualitas operasi layanan", "kecepatan keputusan"],
            },
            selected_chapters=[
                {"title": "BAB I – KONTEKS KLIEN"},
                {"title": "BAB IX – STRUKTUR & TEAM PROYEK"},
            ],
        )

        self.assertNotIn("Konteks akun internal", summary)
        self.assertNotIn("Gunakan informasi ini", summary)
        self.assertNotIn("BAB I", summary)
        self.assertNotIn("BAB IX", summary)
        self.assertNotIn("Problem, Opportunity, Directive", summary)
        self.assertNotIn("kebutuhan Problem, Opportunity, Directive", summary)
        self.assertIn("inisiatif yang diusulkan", summary)
        self.assertIn("kebutuhan prioritas yang perlu dipertegas", summary)

    def test_executive_summary_synthesizes_goal_before_outcome_fallback(self):
        summary = ExecutiveSummaryBuilder.build(
            client="PT Contoh",
            project="Program Transformasi",
            project_goal="Problem, Opportunity, Directive",
            value_map={},
            personalization_pack={},
        )

        self.assertNotIn("Problem, Opportunity, Directive", summary)
        self.assertIn("kebutuhan prioritas yang perlu dipertegas", summary)


if __name__ == "__main__":
    unittest.main()
