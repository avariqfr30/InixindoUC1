"""Regression coverage for the OSINT-backed writer firm profile section."""
from __future__ import annotations

import os
import sys
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))


class WriterFirmProfileEvidenceTest(unittest.TestCase):
    @classmethod
    def setUpClass(cls) -> None:
        cls.tmp = tempfile.TemporaryDirectory()
        tmp_path = Path(cls.tmp.name)
        os.environ.update(
            {
                "APP_STATE_DB_PATH": str(tmp_path / "app_state.db"),
                "APP_ASSET_ROOT": str(tmp_path / "assets"),
                "GENERATED_OUTPUT_DIR": str(tmp_path / "generated"),
                "PROJECT_DB_PATH": str(tmp_path / "projects.db"),
                "PROJECT_CSV_PATH": str(ROOT / "db.csv"),
                "VECTOR_STORE_DIR": str(tmp_path / "chroma"),
                "KB_SYNC_STATE_PATH": str(tmp_path / "kb_state.json"),
                "PROJECT_DATA_SOURCE": "local",
                "APP_PROFILE": "demo",
                "INTERNAL_DATA_SOURCE": "demo",
            }
        )

    @classmethod
    def tearDownClass(cls) -> None:
        cls.tmp.cleanup()

    def test_engine_builds_docx_profile_from_closing_osint_evidence_path(self) -> None:
        from main.proposal_engine import ProposalEngineMixin

        firm_profile = {
            "profile_summary": "Inixindo Jogja membantu organisasi mempercepat kapabilitas digital.",
            "portfolio_highlights": "Portofolio internal training dan konsultasi IT.",
            "credential_highlights": "Instruktur bersertifikasi internasional.",
            "office_address": "Jalan Kenari 69, Yogyakarta",
            "official_source_urls": ["https://inixindojogja.co.id/"],
        }
        osint_evidence = {
            "portfolio_scale": "Lebih dari 1.200 alumni pelatihan profesional tercatat pada kanal publik.",
            "certifications": "Sertifikasi yang disebutkan sumber publik mencakup ISO 27001 dan EC-Council.",
            "team_expertise": "Instruktur publik mencakup praktisi cloud dan cyber security.",
            "accolades": "Profil publik menyebut kolaborasi vendor global.",
        }

        with patch(
            "main.proposal_engine.Researcher.build_comprehensive_firm_profile",
            return_value=osint_evidence,
        ) as build_profile:
            evidence_profile = ProposalEngineMixin._build_writer_firm_evidence_profile(firm_profile)

        build_profile.assert_called_once()
        self.assertIn("1.200 alumni", evidence_profile["portfolio_scale"])
        self.assertIn("ISO 27001", evidence_profile["certifications"])
        self.assertIn("Dirangkum dari sumber publik", evidence_profile["team_expertise"])

    def test_docx_profile_renders_external_credentials_scale_and_source_note(self) -> None:
        from main.document_rendering import DocumentBuilder

        doc = Document()
        profile = {
            "profile_summary": "Inixindo Jogja membantu organisasi mempercepat kapabilitas digital.",
            "portfolio_highlights": "Portofolio internal training dan konsultasi IT.",
            "credential_highlights": "Instruktur bersertifikasi internasional.",
            "portfolio_scale": "Dirangkum dari sumber publik/OSINT: lebih dari 1.200 alumni profesional.",
            "certifications": "Dirangkum dari sumber publik/OSINT: ISO 27001 dan EC-Council.",
            "team_expertise": "Dirangkum dari sumber publik/OSINT: praktisi cloud dan cyber security.",
            "official_source_urls": [
                "https://inixindojogja.co.id/",
                "https://www.inixindo.id/training/it-risk-management/",
            ],
        }

        DocumentBuilder.add_writer_firm_profile_section(doc, profile, (0, 51, 102))

        rendered_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
        rendered_text += "\n" + "\n".join(
            cell.text
            for table in doc.tables
            for row in table.rows
            for cell in row.cells
        )
        self.assertIn("Skala dan bukti portofolio", rendered_text)
        self.assertIn("1.200 alumni", rendered_text)
        self.assertIn("Sertifikasi dan kredensial eksternal", rendered_text)
        self.assertIn("ISO 27001", rendered_text)
        self.assertIn("Keahlian tim berbasis sumber publik", rendered_text)
        self.assertIn("Profil dan kontak pada bagian ini dirangkum", rendered_text)
        self.assertIn("inixindojogja.co.id", rendered_text)

    def test_closing_render_text_strips_inline_writer_profile(self) -> None:
        from main.proposal_engine import ProposalEngineMixin

        closing_text = (
            "Terima kasih atas kesempatan kolaborasi ini.\n\n"
            "## Tentang Mitra Penulis Proposal\n\n"
            "Dirangkum dari sumber publik/OSINT: ISO 27001 dan EC-Council.\n\n"
            "## Informasi Kontak dan Langkah Lanjutan\n\n"
            "- Email: marketing@inixindojogja.co.id"
        )

        cleaned = ProposalEngineMixin._strip_writer_firm_profile_from_closing(closing_text)

        self.assertEqual(cleaned, "Terima kasih atas kesempatan kolaborasi ini.")

    def test_generic_osint_placeholders_are_not_rendered_as_evidence(self) -> None:
        from main.proposal_support import ProposalSupportMixin

        section = ProposalSupportMixin._build_firm_information_section_from_osint(
            firm_name="Inixindo Jogja",
            comprehensive_profile={
                "team_expertise": "Tim profesional Inixindo Jogja memiliki pengalaman di berbagai domain strategis.",
                "portfolio_scale": "Portofolio Inixindo Jogja mencakup berbagai klien enterprise.",
                "certifications": "Kredensial dan sertifikasi Inixindo Jogja tersedia di saluran publik resmi.",
            },
            firm_profile={},
        )

        self.assertEqual(section, "")


if __name__ == "__main__":
    unittest.main()
