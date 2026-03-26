"""Core logic for data access, research, and proposal document generation."""

import os
import io
import re
import json
import logging
import requests
import pandas as pd
import chromadb
from chromadb.config import Settings
import concurrent.futures
from datetime import datetime
from functools import lru_cache
from typing import Dict, List, Any, Tuple, Optional, Set
from urllib.parse import urlparse

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as patches
from PIL import Image, ImageDraw, ImageFont, ImageStat

from sqlalchemy import create_engine
import markdown
from bs4 import BeautifulSoup

from docx import Document
from docx.image.exceptions import UnrecognizedImageError
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from ollama import Client
from chromadb.utils import embedding_functions

from config import (
    SERPER_API_KEY, OLLAMA_HOST, LLM_MODEL, EMBED_MODEL,
    WRITER_FIRM_NAME, DEFAULT_COLOR, UNIVERSAL_STRUCTURE,
    PERSONAS, PROPOSAL_SYSTEM_PROMPT,
    PROJECT_DATA_FIELD_ALIASES, PROJECT_STANDARD_FIELD_ALIASES, FIRM_PROFILE_FIELD_ALIASES,
    CLIENT_RELATIONSHIP_FIELD_ALIASES,
    DEMO_MODE, FIRM_API_URL, API_AUTH_TOKEN, MOCK_FIRM_STANDARDS, MOCK_FIRM_PROFILE,
    DATA_ACQUISITION_MODE,
    WRITER_FIRM_OFFICE_ADDRESS, WRITER_FIRM_EMAIL, WRITER_FIRM_PHONE, WRITER_FIRM_WEBSITE,
    MAX_PROPOSAL_PAGES, ESTIMATED_WORDS_PER_PAGE, RESERVED_NON_CONTENT_PAGES, PAGE_SAFETY_BUFFER,
)

logger = logging.getLogger(__name__)


class SchemaMapper:
    @classmethod
    def flatten_payload(cls, payload: Any, prefix: str = "") -> Dict[str, Any]:
        if not isinstance(payload, dict):
            return {}

        flattened: Dict[str, Any] = {}
        for key, value in payload.items():
            raw_key = str(key or "").strip()
            if not raw_key:
                continue
            next_prefix = f"{prefix}_{raw_key}" if prefix else raw_key
            if isinstance(value, dict):
                flattened.update(cls.flatten_payload(value, next_prefix))
                continue
            if isinstance(value, list) and value and all(not isinstance(item, (dict, list)) for item in value):
                flattened[next_prefix] = ", ".join(str(item) for item in value if str(item).strip())
                continue
            flattened[next_prefix] = value
        return flattened

    @staticmethod
    def normalize_key(value: Any) -> str:
        text = str(value or "").strip().lower()
        text = re.sub(r"[^a-z0-9]+", "_", text)
        return re.sub(r"_+", "_", text).strip("_")

    @classmethod
    def _token_set(cls, value: Any) -> Set[str]:
        return set(filter(None, cls.normalize_key(value).split("_")))

    @classmethod
    def _resolve_alias(cls, raw_key: Any, alias_map: Dict[str, List[str]]) -> Optional[str]:
        normalized_key = cls.normalize_key(raw_key)
        if not normalized_key:
            return None

        direct_hits: List[str] = []
        fuzzy_hits: List[Tuple[int, str]] = []
        raw_tokens = cls._token_set(raw_key)

        for canonical, aliases in (alias_map or {}).items():
            for alias in aliases or []:
                alias_key = cls.normalize_key(alias)
                if not alias_key:
                    continue
                if normalized_key == alias_key:
                    direct_hits.append(canonical)
                    break
                alias_tokens = cls._token_set(alias)
                if raw_tokens and alias_tokens and (alias_tokens <= raw_tokens or raw_tokens <= alias_tokens):
                    fuzzy_hits.append((len(alias_tokens & raw_tokens), canonical))

        if direct_hits:
            return direct_hits[0]
        if fuzzy_hits:
            fuzzy_hits.sort(reverse=True)
            return fuzzy_hits[0][1]
        return None

    @classmethod
    def normalize_record(
        cls,
        payload: Optional[Dict[str, Any]],
        alias_map: Dict[str, List[str]],
        passthrough_keys: Optional[List[str]] = None
    ) -> Dict[str, Any]:
        raw = dict(payload or {})
        flattened = cls.flatten_payload(raw)
        raw.update({k: v for k, v in flattened.items() if k not in raw})
        normalized: Dict[str, Any] = {}
        passthrough = passthrough_keys or []

        for key, value in raw.items():
            mapped_key = cls._resolve_alias(key, alias_map)
            target_key = mapped_key or cls.normalize_key(key)
            if mapped_key and normalized.get(mapped_key) not in (None, "", [], {}):
                continue
            normalized[target_key] = value

        for key in passthrough:
            if key in raw and key not in normalized:
                normalized[key] = raw[key]
        return normalized

    @classmethod
    def remap_dataframe(cls, df: pd.DataFrame, alias_map: Dict[str, List[str]]) -> pd.DataFrame:
        renamed = df.copy()
        rename_map: Dict[str, str] = {}
        taken_targets = set(renamed.columns)

        for column in renamed.columns:
            target = cls._resolve_alias(column, alias_map)
            if not target or column == target or target in taken_targets:
                continue
            rename_map[column] = target
            taken_targets.add(target)

        if rename_map:
            renamed = renamed.rename(columns=rename_map)
        return renamed


# Internal API adapter.
class FirmAPIClient:
    def __init__(self) -> None:
        self.demo_mode = DEMO_MODE
        self.data_acquisition_mode = DATA_ACQUISITION_MODE
        self.base_url = FIRM_API_URL
        self.headers = {"Authorization": f"Bearer {API_AUTH_TOKEN}"}

    def uses_demo_logic(self) -> bool:
        return self.demo_mode or self.data_acquisition_mode == "demo"

    @staticmethod
    def _normalize_project_standards(raw_payload: Optional[Dict[str, Any]]) -> Dict[str, str]:
        normalized = SchemaMapper.normalize_record(raw_payload, PROJECT_STANDARD_FIELD_ALIASES)
        methodology = str(normalized.get("methodology") or "").strip() or "TBD"
        team = str(normalized.get("team") or "").strip() or "TBD"
        commercial = str(normalized.get("commercial") or "").strip() or "TBD"
        return {
            "methodology": methodology,
            "team": team,
            "commercial": commercial,
        }

    @staticmethod
    def _normalize_relationship_mode(value: Any) -> str:
        normalized = SchemaMapper.normalize_key(value)
        existing_markers = {
            "existing", "active", "returning", "repeat", "renewal", "incumbent",
            "prior", "previous", "historical", "yes", "true", "1"
        }
        return "existing" if normalized in existing_markers else "new"

    @classmethod
    def _normalize_client_relationship(cls, raw_payload: Optional[Dict[str, Any]]) -> Dict[str, Any]:
        normalized = SchemaMapper.normalize_record(raw_payload, CLIENT_RELATIONSHIP_FIELD_ALIASES)
        status_value = normalized.get("status", "")
        summary = str(normalized.get("summary") or "").strip()
        mode = cls._normalize_relationship_mode(status_value)
        return {
            "summary": summary,
            "mode": mode,
            "source": "internal_api",
            "verified": bool(summary or str(status_value).strip()),
        }

    @staticmethod
    def _empty_relationship_context(source: str = "internal_api") -> Dict[str, Any]:
        return {
            "summary": "",
            "mode": "new",
            "source": source,
            "verified": False,
        }

    def get_project_standards(self, project_type: str) -> Dict[str, str]:
        if self.demo_mode:
            logger.info("Using demo standards for project type: %s", project_type)
            demo_standards = MOCK_FIRM_STANDARDS.get(project_type, MOCK_FIRM_STANDARDS.get("Implementation"))
            return self._normalize_project_standards(demo_standards)
        try:
            res = requests.get(f"{self.base_url}/standards/{project_type}", headers=self.headers, timeout=5)
            res.raise_for_status()
            return self._normalize_project_standards(res.json())
        except requests.RequestException as e:
            logger.error(f"Internal API Error: {e}")
            return self._normalize_project_standards({})

    @classmethod
    def _default_firm_profile(cls) -> Dict[str, str]:
        base_profile = dict(MOCK_FIRM_PROFILE)
        base_profile["office_address"] = base_profile.get("office_address") or WRITER_FIRM_OFFICE_ADDRESS
        base_profile["email"] = base_profile.get("email") or WRITER_FIRM_EMAIL
        base_profile["phone"] = base_profile.get("phone") or WRITER_FIRM_PHONE
        base_profile["website"] = base_profile.get("website") or WRITER_FIRM_WEBSITE
        return cls._normalize_firm_profile(base_profile)

    def get_firm_profile(self) -> Dict[str, str]:
        if self.demo_mode:
            return self._default_firm_profile()
        try:
            res = requests.get(f"{self.base_url}/firm-profile", headers=self.headers, timeout=5)
            res.raise_for_status()
            return self._normalize_firm_profile(res.json())
        except requests.RequestException as e:
            logger.error(f"Internal API Error: {e}")
            if self.uses_demo_logic():
                return self._build_profile_from_osint()
            return self._default_firm_profile()

    def get_client_relationship(self, client_name: str) -> Dict[str, Any]:
        if self.uses_demo_logic():
            summary = Researcher.get_client_writer_collaboration(client_name, WRITER_FIRM_NAME)
            has_evidence = ProposalGenerator._has_external_evidence(summary)
            return {
                "summary": summary,
                "mode": "existing" if has_evidence else "new",
                "source": "osint",
                "verified": has_evidence,
            }

        try:
            res = requests.get(
                f"{self.base_url}/client-relationship",
                params={"client_name": client_name},
                headers=self.headers,
                timeout=5
            )
            res.raise_for_status()
            return self._normalize_client_relationship(res.json())
        except requests.RequestException as e:
            logger.error(f"Internal API Error: {e}")
            return self._empty_relationship_context()

    @staticmethod
    def _extract_first(pattern: str, text: str) -> str:
        match = re.search(pattern, text or "", flags=re.IGNORECASE)
        return match.group(0).strip() if match else ""

    @staticmethod
    def _clean_contact_value(value: str) -> str:
        cleaned = re.sub(r"\s+", " ", (value or "").replace("\xa0", " ")).strip()
        return cleaned.strip(" ,;|-")

    @classmethod
    def _looks_like_address(cls, value: str) -> bool:
        cleaned = cls._clean_contact_value(value)
        lowered = cleaned.lower()
        if len(cleaned) < 14:
            return False
        if "@" in cleaned or re.search(r"https?://|www\.", lowered):
            return False
        address_markers = (
            "jl.", "jalan", "alamat", "gedung", "tower", "ruko", "komplek", "kompleks",
            "kawasan", "suite", "lantai", "yogyakarta", "sleman", "bantul", "jakarta",
            "bandung", "surabaya", "indonesia"
        )
        return any(marker in lowered for marker in address_markers)

    @classmethod
    def _extract_address_candidates(cls, text: str) -> List[str]:
        if not text:
            return []
        candidates: List[str] = []
        patterns = [
            r"(?:alamat|address)\s*[:\-]?\s*([^|\n]{12,180})",
            r"((?:jl\.|jalan)\s+[^|\n]{12,180})",
        ]
        for pattern in patterns:
            for match in re.finditer(pattern, text, flags=re.IGNORECASE):
                candidate = match.group(1) if match.lastindex else match.group(0)
                candidate = re.split(r"(?:\s+[|]\s+| email\s*:| telp\s*:| phone\s*:| website\s*:)", candidate, maxsplit=1, flags=re.IGNORECASE)[0]
                candidate = cls._clean_contact_value(candidate)
                if cls._looks_like_address(candidate):
                    candidates.append(candidate)

        unique_candidates: List[str] = []
        seen = set()
        for candidate in candidates:
            key = candidate.lower()
            if key in seen:
                continue
            unique_candidates.append(candidate)
            seen.add(key)
        return unique_candidates

    @classmethod
    def _extract_contact_fields(cls, text: str) -> Dict[str, str]:
        office_candidates = cls._extract_address_candidates(text)
        website = cls._extract_first(r"(?:https?://|www\.)[A-Za-z0-9./_%#?=&-]+\.[A-Za-z]{2,}", text)
        return {
            "office_address": office_candidates[0] if office_candidates else "",
            "email": cls._extract_first(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", text),
            "phone": cls._extract_first(r"(?:\+62|62|0)\d[\d\-\s()]{7,}\d", text),
            "website": website.rstrip(".,;)"),
        }

    @classmethod
    def build_contact_lines(cls, firm_profile: Optional[Dict[str, Any]]) -> List[str]:
        profile = firm_profile or {}
        lines = []
        office_address = cls._clean_contact_value(str(profile.get("office_address") or ""))
        email = cls._clean_contact_value(str(profile.get("email") or ""))
        phone = cls._clean_contact_value(str(profile.get("phone") or ""))
        website = cls._clean_contact_value(str(profile.get("website") or ""))

        if office_address:
            lines.append(f"Alamat kantor: {office_address}")
        if email:
            lines.append(f"Email: {email}")
        if phone:
            lines.append(f"Telp: {phone}")
        if website:
            if not re.match(r"^https?://", website, flags=re.IGNORECASE):
                website = f"https://{website.lstrip('/')}"
            lines.append(f"Website: {website}")
        return lines

    @classmethod
    def _normalize_firm_profile(cls, raw_profile: Optional[Dict[str, Any]]) -> Dict[str, str]:
        source = SchemaMapper.normalize_record(raw_profile, FIRM_PROFILE_FIELD_ALIASES)
        raw_contact = str(source.get("contact_info") or "").strip()
        parsed = cls._extract_contact_fields(raw_contact)

        office_address = cls._clean_contact_value(str(
            source.get("office_address")
            or source.get("address")
            or source.get("office_location")
            or source.get("location")
            or parsed.get("office_address")
            or ""
        ))
        email = cls._clean_contact_value(str(
            source.get("email")
            or source.get("official_email")
            or parsed.get("email")
            or ""
        ))
        phone = cls._clean_contact_value(str(
            source.get("phone")
            or source.get("telephone")
            or source.get("telp")
            or source.get("mobile")
            or parsed.get("phone")
            or ""
        ))
        website = cls._clean_contact_value(str(
            source.get("website")
            or source.get("url")
            or parsed.get("website")
            or ""
        ))
        portfolio = cls._clean_contact_value(str(source.get("portfolio_highlights") or ""))
        contact_lines = cls.build_contact_lines(
            {
                "office_address": office_address,
                "email": email,
                "phone": phone,
                "website": website,
            }
        )

        return {
            "office_address": office_address,
            "email": email,
            "phone": phone,
            "website": website,
            "contact_info": "\n".join(contact_lines),
            "portfolio_highlights": portfolio or "Kapabilitas layanan menyesuaikan kebutuhan proyek klien.",
        }

    def _build_profile_from_osint(self) -> Dict[str, str]:
        current_year = datetime.now().year
        query = (
            f'"{WRITER_FIRM_NAME}" Yogyakarta kontak email telp alamat profil '
            f'{current_year} OR {current_year - 1}'
        )
        raw_hits = Researcher.search(query, limit=8, recency_bucket="year")
        hits = Researcher._filter_recent_entity_results(
            raw_hits,
            entity_name=WRITER_FIRM_NAME,
            max_age_years=4,
            strict_entity=True
        )
        if not hits:
            fallback_query = f'"{WRITER_FIRM_NAME}" Yogyakarta pelatihan konsultasi kontak resmi'
            fallback_hits = Researcher.search(fallback_query, limit=8)
            hits = Researcher._filter_recent_entity_results(
                fallback_hits,
                entity_name=WRITER_FIRM_NAME,
                max_age_years=6,
                strict_entity=True
            )
        address_query = f'"{WRITER_FIRM_NAME}" alamat kantor Yogyakarta OR "Jl." OR "Jalan"'
        address_hits = Researcher.search(address_query, limit=8, recency_bucket="year")
        address_hits = Researcher._filter_recent_entity_results(
            address_hits,
            entity_name=WRITER_FIRM_NAME,
            max_age_years=6,
            strict_entity=True
        )

        merged_hits = hits + [item for item in address_hits if item not in hits]
        merged_text = " ".join(
            [
                " ".join(
                    [
                        str(item.get("title", "")),
                        str(item.get("snippet", "")),
                        str(item.get("link", "")),
                    ]
                )
                for item in merged_hits
            ]
        )
        parsed = self._extract_contact_fields(merged_text)

        if not parsed.get("office_address"):
            for item in merged_hits:
                item_text = " ".join(
                    [
                        str(item.get("title", "")),
                        str(item.get("snippet", "")),
                        str(item.get("link", "")),
                    ]
                )
                parsed_item = self._extract_contact_fields(item_text)
                if parsed_item.get("office_address"):
                    parsed["office_address"] = parsed_item["office_address"]
                    break

        return self._normalize_firm_profile(
            {
                "office_address": parsed.get("office_address", ""),
                "email": parsed.get("email", ""),
                "phone": parsed.get("phone", ""),
                "website": parsed.get("website", ""),
                "portfolio_highlights": "Kapabilitas layanan menyesuaikan kebutuhan proyek klien.",
            }
        )


# Knowledge base and vector index.
class KnowledgeBase:
    def __init__(self, db_uri: str) -> None:
        self.engine = create_engine(db_uri)
        self.chroma = chromadb.Client(Settings(anonymized_telemetry=False))
        self.embed_fn = embedding_functions.OllamaEmbeddingFunction(
            url=f"{OLLAMA_HOST}/api/embeddings", model_name=EMBED_MODEL
        )
        self.collection = self.chroma.get_or_create_collection(
            name="projects_db", embedding_function=self.embed_fn
        )
        self.df: Optional[pd.DataFrame] = None
        self.refresh_data()

    @staticmethod
    def _required_project_fields() -> Tuple[str, ...]:
        return ("entity", "topic")

    @classmethod
    def _normalize_projects_df(cls, raw_df: pd.DataFrame) -> pd.DataFrame:
        normalized = raw_df.copy()
        normalized.columns = [str(col).strip() for col in normalized.columns]
        normalized = SchemaMapper.remap_dataframe(normalized, PROJECT_DATA_FIELD_ALIASES)
        return normalized

    @classmethod
    def _has_required_project_fields(cls, df: Optional[pd.DataFrame]) -> bool:
        if df is None:
            return False
        return all(field in df.columns for field in cls._required_project_fields())

    def refresh_data(self) -> bool:
        try:
            self.df = self._normalize_projects_df(pd.read_sql("SELECT * FROM projects", self.engine))
        except Exception:
            if not os.path.exists("db.csv"):
                return False
            raw_df = pd.read_csv("db.csv")
            normalized_df = self._normalize_projects_df(raw_df)
            normalized_df.to_sql("projects", self.engine, index=False, if_exists='replace')
            self.df = normalized_df

        if not self._has_required_project_fields(self.df):
            logger.warning(
                "Project data schema is missing required fields. Expected aliases for: %s. Available columns: %s",
                ", ".join(self._required_project_fields()),
                ", ".join(self.df.columns.astype(str).tolist()) if self.df is not None else "-",
            )
            return False
            
        existing_ids = set(self.collection.get()['ids'])
        new_ids_map = {str(idx): row for idx, row in self.df.iterrows()}
        new_ids_set = set(new_ids_map.keys())
        
        ids_to_delete = list(existing_ids - new_ids_set)
        
        if ids_to_delete: 
            self.collection.delete(ids_to_delete)

        all_ids = list(new_ids_set)
        if all_ids:
            for i in range(0, len(all_ids), 500):
                batch_ids = all_ids[i:i + 500]
                docs = [" | ".join([f"{col}: {val}" for col, val in new_ids_map[b].items()]) for b in batch_ids]
                metas = [new_ids_map[b].astype(str).to_dict() for b in batch_ids]
                self.collection.upsert(documents=docs, metadatas=metas, ids=batch_ids)
                
        return True

    def get_exact_context(self, entity: str, topic: str, budget: Optional[str] = None) -> str:
        if self.df is None or self.df.empty or not self._has_required_project_fields(self.df):
            return "No data."
        try:
            match = self.df[(self.df['entity'] == entity) & (self.df['topic'] == topic)]
            if budget and 'budget' in self.df.columns and not match.empty:
                match = match[match['budget'] == budget]
            if not match.empty:
                return "".join([f"- {k.capitalize()}: {v}\n" for k, v in match.iloc[0].to_dict().items()])
            return "No data."
        except Exception:
            return ""

    def query(self, client: str, project: str, context_keywords: str = "") -> str:
        try:
            res = self.collection.query(query_texts=[f"{project} for {client} {context_keywords}"], n_results=2)
            if res['documents'] and len(res['documents'][0]) > 0:
                return "\n".join(res['documents'][0])
            return ""
        except Exception:
            return ""


# Public research helper (Serper).
class Researcher:
    @staticmethod
    def _has_serper_key() -> bool:
        key = (SERPER_API_KEY or "").strip()
        if not key:
            return False
        placeholder_keys = {"YOUR_SERPER_API_KEY", "YOUR_SERPER", "SERPER_API"}
        return key not in placeholder_keys

    @staticmethod
    @lru_cache(maxsize=256)
    def search(query: str, limit: int = 5, recency_bucket: str = "") -> List[Dict[str, Any]]:
        """General web search using Serper.dev"""
        if not Researcher._has_serper_key():
            return []
        url = "https://google.serper.dev/search"
        payload_data = {"q": query, "gl": "id", "num": limit}
        recency_map = {
            "week": "qdr:w",
            "month": "qdr:m",
            "year": "qdr:y",
        }
        if recency_bucket in recency_map:
            payload_data["tbs"] = recency_map[recency_bucket]
        payload = json.dumps(payload_data)
        headers = {
            'X-API-KEY': (SERPER_API_KEY or "").strip(),
            'Content-Type': 'application/json'
        }
        try:
            response = requests.post(url, headers=headers, data=payload, timeout=8)
            response.raise_for_status()
            return response.json().get('organic', [])
        except requests.RequestException as e:
            logger.warning(f"Serper API Error: {e}")
            return []

    @staticmethod
    def _normalize_text(text: str) -> str:
        normalized = re.sub(r"[^a-z0-9]+", " ", (text or "").lower())
        return re.sub(r"\s+", " ", normalized).strip()

    @staticmethod
    def _entity_tokens(entity_name: str) -> List[str]:
        legal_tokens = {"pt", "cv", "tbk", "inc", "ltd", "co", "corp", "persero", "company"}
        tokens = [
            token for token in re.findall(r"[a-z0-9]+", (entity_name or "").lower())
            if len(token) >= 3 and token not in legal_tokens
        ]
        # Keep order while dropping duplicates.
        ordered = []
        seen = set()
        for token in tokens:
            if token in seen:
                continue
            ordered.append(token)
            seen.add(token)
        return ordered

    @staticmethod
    def _is_entity_match(item: Dict[str, Any], entity_name: str, strict: bool = False) -> bool:
        if not entity_name:
            return True

        merged = " ".join([
            str(item.get("title", "")),
            str(item.get("snippet", "")),
            str(item.get("link", "")),
        ])
        normalized_merged = Researcher._normalize_text(merged)
        phrase = Researcher._normalize_text(entity_name)
        tokens = Researcher._entity_tokens(entity_name)

        if phrase and phrase in normalized_merged:
            return True

        if not tokens:
            return False

        merged_tokens = set(normalized_merged.split())
        hits = sum(1 for token in tokens if token in merged_tokens)
        if strict:
            return hits == len(tokens)
        if len(tokens) == 1:
            return hits == 1
        if len(tokens) == 2:
            return hits == 2
        return hits >= (len(tokens) - 1)

    @staticmethod
    def _extract_month(text: str) -> Optional[int]:
        if not text:
            return None
        month_map = {
            "jan": 1, "januari": 1, "january": 1,
            "feb": 2, "februari": 2, "february": 2,
            "mar": 3, "maret": 3, "march": 3,
            "apr": 4, "april": 4,
            "mei": 5, "may": 5,
            "jun": 6, "juni": 6, "june": 6,
            "jul": 7, "juli": 7, "july": 7,
            "agu": 8, "agustus": 8, "aug": 8, "august": 8,
            "sep": 9, "sept": 9, "september": 9,
            "okt": 10, "oct": 10, "oktober": 10, "october": 10,
            "nov": 11, "november": 11,
            "des": 12, "dec": 12, "desember": 12, "december": 12,
        }
        lowered = (text or "").lower()
        for key, month in month_map.items():
            if re.search(rf"\b{re.escape(key)}\b", lowered):
                return month
        return None

    @staticmethod
    def _extract_day(text: str) -> Optional[int]:
        if not text:
            return None
        match = re.search(r"\b([0-2]?\d|3[01])\b", text)
        if not match:
            return None
        day = int(match.group(1))
        if 1 <= day <= 31:
            return day
        return None

    @staticmethod
    def _published_sort_key(item: Dict[str, Any]) -> int:
        merged = " ".join([
            str(item.get("date", "")),
            str(item.get("title", "")),
            str(item.get("snippet", "")),
        ])
        year = Researcher._extract_year(merged) or 0
        month = Researcher._extract_month(merged) or 1
        day = Researcher._extract_day(str(item.get("date", ""))) or 1
        return (year * 10_000) + (month * 100) + day

    @staticmethod
    def _sort_by_recency(items: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        return sorted(items or [], key=Researcher._published_sort_key, reverse=True)

    @staticmethod
    def _filter_recent_entity_results(
        items: List[Dict[str, Any]],
        entity_name: str = "",
        max_age_years: int = 2,
        strict_entity: bool = False
    ) -> List[Dict[str, Any]]:
        filtered = []
        for item in (items or []):
            if not Researcher._is_recent(item, max_age_years=max_age_years):
                continue
            if entity_name and not Researcher._is_entity_match(item, entity_name, strict=strict_entity):
                continue
            filtered.append(item)
        return Researcher._sort_by_recency(filtered)

    @staticmethod
    def _is_regulatory_domain(link: str) -> bool:
        domain = Researcher._source_name(link)
        trusted_suffixes = ("go.id", "iso.org", "ietf.org", "iec.ch", "nist.gov")
        return any(domain.endswith(sfx) for sfx in trusted_suffixes)
    @staticmethod
    def _extract_year(text: str) -> Optional[int]:
        if not text:
            return None
        years = re.findall(r'\b(20\d{2})\b', text)
        if not years:
            return None
        return max(int(y) for y in years)

    @staticmethod
    def _is_recent(item: Dict[str, Any], max_age_years: int = 2) -> bool:
        merged = " ".join([
            str(item.get('date', '')),
            str(item.get('snippet', '')),
            str(item.get('title', '')),
        ])
        year = Researcher._extract_year(merged)
        if year is None:
            return True
        return year >= (datetime.now().year - max_age_years)

    @staticmethod
    def _source_name(link: str) -> str:
        try:
            host = urlparse(link or "").netloc.lower().strip()
            host = host.replace("www.", "")
            return host or "sumber daring"
        except Exception:
            return "sumber daring"

    @staticmethod
    def _citation_year(item: Dict[str, Any]) -> str:
        merged = " ".join([
            str(item.get('date', '')),
            str(item.get('snippet', '')),
            str(item.get('title', '')),
        ])
        year = Researcher._extract_year(merged)
        return str(year) if year else "n.d."

    @staticmethod
    def _format_evidence(items: List[Dict[str, Any]], label: str, fallback: str) -> str:
        if not items:
            return f"{fallback} (sumber daring, n.d.)"
        lines = []
        for i, item in enumerate(items, start=1):
            title = item.get('title', 'Sumber tanpa judul')
            snippet = (item.get('snippet', '') or '').strip()
            link = item.get('link', '-')
            if not snippet:
                continue
            source_name = Researcher._source_name(link)
            citation = f"({source_name}, {Researcher._citation_year(item)})"
            lines.append(
                f"Sumber eksternal {i}: fakta={snippet} | sumber={title} | url={link} | sitasi_apa={citation}"
            )
        return "\n".join(lines) if lines else f"{fallback} (sumber daring, n.d.)"


    @staticmethod
    @lru_cache(maxsize=128)
    def get_entity_profile(entity_name: str) -> str:
        current_year = datetime.now().year
        strict_entity = Researcher._normalize_text(entity_name) == Researcher._normalize_text(WRITER_FIRM_NAME)
        query = (
            f'"{entity_name}" profil perusahaan OR "tentang kami" OR alamat OR kontak '
            f'{current_year} OR {current_year - 1}'
        )
        res = Researcher.search(query, limit=8, recency_bucket="year")
        filtered = Researcher._filter_recent_entity_results(
            res,
            entity_name=entity_name,
            max_age_years=3,
            strict_entity=strict_entity
        )
        if not filtered and strict_entity:
            fallback_query = f'"{entity_name}" Yogyakarta pelatihan konsultasi kontak resmi'
            fallback_res = Researcher.search(fallback_query, limit=8)
            filtered = Researcher._filter_recent_entity_results(
                fallback_res,
                entity_name=entity_name,
                max_age_years=6,
                strict_entity=True
            )
        return Researcher._format_evidence(
            filtered[:4],
            label="OSINT_PROFILE",
            fallback=f"Data profil terbaru untuk {entity_name} terbatas; gunakan informasi umum yang terverifikasi saja."
        )

    @staticmethod
    @lru_cache(maxsize=128)
    def get_latest_client_news(client_name: str) -> str:
        current_year = datetime.now().year
        prev_year = current_year - 1
        res = Researcher.search(
            f'"{client_name}" berita inovasi OR transformasi digital {current_year} OR {prev_year}',
            limit=8,
            recency_bucket="month"
        )
        filtered = Researcher._filter_recent_entity_results(
            res,
            entity_name=client_name,
            max_age_years=2,
            strict_entity=False
        )
        return Researcher._format_evidence(
            filtered[:4],
            label="OSINT_NEWS",
            fallback=f"Berita terbaru {client_name} tidak cukup kuat; jangan membuat klaim spesifik tanpa bukti."
        )

    @staticmethod
    @lru_cache(maxsize=128)
    def get_client_track_record(client_name: str) -> str:
        current_year = datetime.now().year
        prev_year = current_year - 1
        res = Researcher.search(
            (
                f'"{client_name}" pencapaian OR kinerja OR ekspansi OR transformasi OR penghargaan '
                f'{current_year} OR {prev_year}'
            ),
            limit=10,
            recency_bucket="year"
        )
        filtered = Researcher._filter_recent_entity_results(
            res,
            entity_name=client_name,
            max_age_years=3,
            strict_entity=False
        )
        return Researcher._format_evidence(
            filtered[:5],
            label="OSINT_TRACK",
            fallback=f"Track record publik {client_name} terbatas; hindari klaim performa tanpa bukti."
        )

    @staticmethod
    @lru_cache(maxsize=128)
    def get_client_writer_collaboration(client_name: str, writer_firm_name: str = WRITER_FIRM_NAME) -> str:
        current_year = datetime.now().year
        strict_writer = Researcher._normalize_text(writer_firm_name) == Researcher._normalize_text(WRITER_FIRM_NAME)
        query = (
            f'"{writer_firm_name}" "{client_name}" kerja sama OR proyek OR pelatihan OR konsultasi '
            f'{current_year} OR {current_year - 1} OR {current_year - 2}'
        )
        res = Researcher.search(query, limit=10, recency_bucket="year")
        filtered: List[Dict[str, Any]] = []
        for item in Researcher._sort_by_recency(res):
            if not Researcher._is_recent(item, max_age_years=6):
                continue
            if not Researcher._is_entity_match(item, client_name, strict=False):
                continue
            if not Researcher._is_entity_match(item, writer_firm_name, strict=strict_writer):
                continue
            filtered.append(item)

        return Researcher._format_evidence(
            filtered[:4],
            label="OSINT_COLLAB",
            fallback=(
                f"Belum ditemukan bukti publik yang cukup kuat terkait kolaborasi {writer_firm_name} "
                f"dengan {client_name}; jangan menyatakan pernah bekerja sama tanpa verifikasi."
            )
        )

    @staticmethod
    @lru_cache(maxsize=128)
    def get_regulatory_data(regulations_string: str) -> str:
        if not regulations_string:
            return "Tidak ada regulasi spesifik dari input user."

        query = f'Ringkasan implementasi standar {regulations_string.replace(",", " OR ")} site:.go.id OR site:iso.org'
        res = Researcher.search(query, limit=8, recency_bucket="year")
        recent = [i for i in Researcher._sort_by_recency(res) if Researcher._is_recent(i, max_age_years=5)]
        trusted = [i for i in recent if Researcher._is_regulatory_domain(str(i.get("link", "")))]
        filtered = trusted if trusted else recent
        return Researcher._format_evidence(
            filtered[:5],
            label="OSINT_REG",
            fallback=f"Data regulasi untuk {regulations_string} terbatas; nyatakan asumsi dan batasan data secara eksplisit."
        )


# Budget estimator from public financial context.
class FinancialAnalyzer:
    def __init__(self, ollama_client: Client):
        self.ollama = ollama_client

    @staticmethod
    def _format_idr(amount: int) -> str:
        amount = max(0, int(amount))
        return "Rp " + f"{amount:,}".replace(",", ".")

    @staticmethod
    def _parse_number(raw: str) -> Optional[float]:
        value = (raw or "").strip()
        if not value:
            return None
        if "." in value and "," in value:
            value = value.replace(".", "").replace(",", ".")
        elif "," in value and "." not in value:
            value = value.replace(",", ".")
        elif "." in value and re.fullmatch(r"\d{1,3}(?:\.\d{3})+", value):
            value = value.replace(".", "")
        try:
            return float(value)
        except ValueError:
            return None

    @classmethod
    def _extract_financial_values(cls, text: str) -> List[int]:
        if not text:
            return []
        pattern = re.compile(
            r"(?i)(?P<rp>rp\.?\s*)?(?P<num>\d{1,3}(?:[.,]\d{3})+|\d+(?:[.,]\d+)?)\s*(?P<unit>triliun|miliar|juta|ribu)?"
        )
        multiplier = {
            "triliun": 1_000_000_000_000,
            "miliar": 1_000_000_000,
            "juta": 1_000_000,
            "ribu": 1_000,
        }
        values: List[int] = []
        for match in pattern.finditer(text):
            has_rp = bool(match.group("rp"))
            unit = (match.group("unit") or "").lower()
            base = cls._parse_number(match.group("num") or "")
            if base is None:
                continue
            if not has_rp and not unit:
                continue
            amount = int(base * multiplier.get(unit, 1))
            if amount <= 0:
                continue
            if amount < 1_000_000 and not unit:
                continue
            values.append(amount)
        return values

    @classmethod
    def _duration_to_months(cls, timeline: str) -> Optional[float]:
        text = (timeline or "").lower().strip()
        if not text:
            return None

        patterns = [
            (r"(\d+(?:[.,]\d+)?)\s*(tahun|thn|year|years)", 12.0),
            (r"(\d+(?:[.,]\d+)?)\s*(bulan|bln|month|months)", 1.0),
            (r"(\d+(?:[.,]\d+)?)\s*(minggu|week|weeks)", 1.0 / 4.345),
            (r"(\d+(?:[.,]\d+)?)\s*(hari|day|days)", 1.0 / 30.0),
        ]

        months = 0.0
        found = False
        for pattern, factor in patterns:
            for match in re.finditer(pattern, text):
                value = cls._parse_number(match.group(1))
                if value is None:
                    continue
                months += value * factor
                found = True

        if found:
            return max(months, 0.5)

        single_number = re.search(r"(\d+(?:[.,]\d+)?)", text)
        if single_number:
            value = cls._parse_number(single_number.group(1))
            if value is not None:
                return max(value, 0.5)
        return None

    @classmethod
    def _duration_multiplier(cls, timeline: str) -> float:
        months = cls._duration_to_months(timeline)
        if months is None:
            return 1.0
        raw = (months / 6.0) ** 0.45
        return max(0.75, min(2.4, raw))

    @classmethod
    def _scope_multiplier(
        cls,
        project_type: str,
        service_type: str,
        project_goal: str,
        objective: str,
        notes: str,
        frameworks: str
    ) -> float:
        project_weight = {
            "diagnostic": 0.90,
            "strategic": 1.00,
            "transformation": 1.22,
            "implementation": 1.35,
        }
        service_weight = {
            "training": 0.85,
            "konsultan": 1.10,
            "training dan konsultan": 1.20,
        }

        base = project_weight.get((project_type or "").strip().lower(), 1.0)
        base *= service_weight.get((service_type or "").strip().lower(), 1.0)

        combined = " ".join([project_goal or "", objective or "", notes or "", frameworks or ""]).lower()
        high_complexity = [
            "multi", "integrasi", "core banking", "migrasi", "nasional", "enterprise",
            "regulasi", "compliance", "24/7", "high availability", "multi-site", "multisite",
        ]
        medium_complexity = [
            "dashboard", "governance", "kpi", "workflow", "automation",
            "cloud", "api", "security", "audit", "change management",
        ]
        high_hits = sum(1 for token in high_complexity if token in combined)
        medium_hits = sum(1 for token in medium_complexity if token in combined)
        complexity_boost = min(0.55, (high_hits * 0.06) + (medium_hits * 0.025))

        need_items = [part.strip() for part in re.split(r"[,+;/]| dan ", (project_goal or "").lower()) if part.strip()]
        breadth_boost = min(0.25, max(0, len(need_items) - 1) * 0.05)

        scope_multiplier = base * (1.0 + complexity_boost + breadth_boost)
        return max(0.75, min(2.4, scope_multiplier))

    @staticmethod
    def _compact_keywords(text: str, max_terms: int = 7) -> str:
        stopwords = {
            "dan", "yang", "untuk", "dengan", "dari", "pada", "agar", "atau",
            "the", "and", "for", "with", "from", "into", "this", "that",
        }
        tokens = re.findall(r"[A-Za-z0-9]{4,}", (text or "").lower())
        seen = set()
        out = []
        for token in tokens:
            if token in stopwords or token in seen:
                continue
            seen.add(token)
            out.append(token)
            if len(out) >= max_terms:
                break
        return " ".join(out)

    @classmethod
    def _dynamic_budget_from_osint(
        cls,
        client_name: str,
        finance_snippets: List[Dict[str, Any]],
        benchmark_snippets: List[Dict[str, Any]],
        timeline: str = "",
        project_type: str = "",
        service_type: str = "",
        project_goal: str = "",
        objective: str = "",
        notes: str = "",
        frameworks: str = "",
    ) -> Dict[str, Any]:
        finance_text = " ".join(
            [str(item.get("title", "")) + " " + str(item.get("snippet", "")) for item in (finance_snippets or [])]
        )
        benchmark_text = " ".join(
            [str(item.get("title", "")) + " " + str(item.get("snippet", "")) for item in (benchmark_snippets or [])]
        )

        finance_values = sorted(cls._extract_financial_values(finance_text))
        benchmark_values = sorted(cls._extract_financial_values(benchmark_text))

        if finance_values:
            finance_median = finance_values[len(finance_values) // 2]
            financial_base = int(max(120_000_000, min(6_000_000_000, finance_median * 0.0015)))
        else:
            finance_median = None
            financial_base = 350_000_000

        if benchmark_values:
            benchmark_median = benchmark_values[len(benchmark_values) // 2]
            market_base = int(max(80_000_000, min(6_000_000_000, benchmark_median)))
        else:
            benchmark_median = None
            market_base = financial_base

        if finance_median and benchmark_median:
            base_price = int((financial_base * 0.45) + (market_base * 0.55))
        elif benchmark_median:
            base_price = market_base
        else:
            base_price = financial_base

        duration_factor = cls._duration_multiplier(timeline)
        scope_factor = cls._scope_multiplier(
            project_type=project_type,
            service_type=service_type,
            project_goal=project_goal,
            objective=objective,
            notes=notes,
            frameworks=frameworks,
        )
        adjusted_base = int(base_price * duration_factor * scope_factor)
        adjusted_base = int(max(120_000_000, min(9_000_000_000, adjusted_base)))

        basic = int(adjusted_base * 0.72)
        standard = max(basic + 40_000_000, adjusted_base)
        enterprise = max(standard + 80_000_000, int(adjusted_base * 1.65))

        months = cls._duration_to_months(timeline)
        duration_note = f"{months:.1f} bulan" if months else "durasi belum spesifik"
        if finance_median:
            analysis = (
                f"Estimasi untuk {client_name} memakai sinyal finansial publik "
                f"(acuan {cls._format_idr(finance_median)}) lalu disesuaikan oleh durasi ({duration_note}) "
                f"dan skala/scope proyek."
            )
        else:
            analysis = (
                f"Data finansial publik {client_name} terbatas; estimasi dibuat dari benchmark OSINT "
                f"dan disesuaikan oleh durasi ({duration_note}) serta skala/scope proyek."
            )

        return {
            "analysis": analysis,
            "options": [
                {"tier": "Basic", "price": cls._format_idr(basic)},
                {"tier": "Standard", "price": cls._format_idr(standard)},
                {"tier": "Enterprise", "price": cls._format_idr(enterprise)},
            ],
        }

    @staticmethod
    def _is_valid_budget_payload(payload: Dict[str, Any]) -> bool:
        if not isinstance(payload, dict):
            return False
        options = payload.get("options")
        analysis = payload.get("analysis")
        if not isinstance(analysis, str) or not analysis.strip():
            return False
        if not isinstance(options, list) or len(options) < 3:
            return False
        for item in options[:3]:
            if not isinstance(item, dict):
                return False
            tier = str(item.get("tier", "")).strip()
            price = str(item.get("price", "")).strip()
            if not tier or not price:
                return False
            if "<" in price or ">" in price:
                return False
            if not re.search(r"\d", price):
                return False
        return True

    @staticmethod
    def _merge_with_context_adjustment(
        model_payload: Dict[str, Any],
        context_payload: Dict[str, Any],
        context_sensitive: bool,
        prefer_context_options: bool = True
    ) -> Dict[str, Any]:
        if not context_sensitive:
            return model_payload
        analysis = str(model_payload.get("analysis", "")).strip()
        if analysis:
            analysis = (
                f"{analysis} Rentang harga sudah disesuaikan lagi dengan durasi dan skala proyek dari input."
            )
        else:
            analysis = context_payload.get("analysis", "")
        if not prefer_context_options:
            merged = dict(model_payload)
            merged["analysis"] = analysis
            return merged
        return {
            "analysis": analysis,
            "options": context_payload.get("options", []),
        }

    def suggest_budget(
        self,
        client_name: str,
        timeline: str = "",
        project_type: str = "",
        service_type: str = "",
        project_goal: str = "",
        objective: str = "",
        notes: str = "",
        frameworks: str = "",
        commercial_context: str = "",
        pricing_mode: str = "demo",
    ) -> Dict[str, Any]:
        year = datetime.now().year
        finance_results = Researcher.search(
            f'"{client_name}" laporan keuangan OR pendapatan OR pendanaan OR aset {year-2} OR {year-1} OR {year}',
            limit=10,
            recency_bucket="year"
        )
        finance_snippets = Researcher._filter_recent_entity_results(
            finance_results,
            entity_name=client_name,
            max_age_years=3,
            strict_entity=False
        )
        keyword_context = self._compact_keywords(
            f"{objective} {notes} {project_goal} {frameworks} {project_type} {service_type}"
        )
        benchmark_query = (
            f'estimasi biaya proyek {project_type or "IT"} {service_type} {timeline} '
            f'Indonesia {keyword_context}'
        )
        benchmark_results = Researcher.search(benchmark_query, limit=10, recency_bucket="year")
        benchmark_snippets = [
            item for item in Researcher._sort_by_recency(benchmark_results)
            if Researcher._is_recent(item, max_age_years=3)
        ][:8]

        context_finance = "\n".join([item.get('snippet', '') for item in finance_snippets]) if finance_snippets else "-"
        context_benchmark = "\n".join([item.get('snippet', '') for item in benchmark_snippets]) if benchmark_snippets else "-"
        dynamic_estimate = self._dynamic_budget_from_osint(
            client_name=client_name,
            finance_snippets=finance_snippets,
            benchmark_snippets=benchmark_snippets,
            timeline=timeline,
            project_type=project_type,
            service_type=service_type,
            project_goal=project_goal,
            objective=objective,
            notes=notes,
            frameworks=frameworks,
        )
        staged_pricing = pricing_mode == "staged"
        commercial_note = (
            f"Baseline commercial rules internal:\n{commercial_context}\n"
            "Jadikan baseline internal ini sebagai acuan utama penentuan harga. "
            "Data OSINT hanya dipakai sebagai sinyal pendukung.\n"
            if staged_pricing and commercial_context.strip()
            else ""
        )

        prompt = f"""
        Menganalisa kekuatan finansial perusahaan: {client_name}.
        Data finansial OSINT:
        {context_finance}

        Data benchmark biaya OSINT:
        {context_benchmark}

        Konteks proyek:
        - Durasi: {timeline or '-'}
        - Jenis Proyek: {project_type or '-'}
        - Jenis Proposal/Layanan: {service_type or '-'}
        - Klasifikasi Kebutuhan: {project_goal or '-'}
        - Objective: {objective or '-'}
        - Pain Points: {notes or '-'}
        - Framework: {frameworks or '-'}
        {commercial_note}

        Berdasarkan data di atas, estimasikan kapasitas finansial mereka dan berikan 3 opsi estimasi budget proyek TI/Konsultasi.
        FORMAT WAJIB JSON murni tanpa markdown, tanpa teks tambahan:
        {{
            "analysis": "Ringkasan 1 kalimat kekuatan finansial berdasarkan data (atau sebutkan estimasi jika data terbatas).",
            "options": [
                {{"tier": "Basic", "price": "Rp <angka>"}},
                {{"tier": "Standard", "price": "Rp <angka>"}},
                {{"tier": "Enterprise", "price": "Rp <angka>"}}
            ]
        }}
        Pastikan angka mempertimbangkan durasi dan skala/scope proyek selain data OSINT.
        """
        try:
            res = self.ollama.chat(
                model=LLM_MODEL,
                messages=[{'role': 'system', 'content': 'You output strictly valid JSON.'}, {'role': 'user', 'content': prompt}],
                options={'temperature': 0.1}
            )
            raw_text = res['message']['content']
            match = re.search(r'\{.*\}', raw_text, re.DOTALL)
            parsed = {}
            if match:
                parsed = json.loads(match.group(0))
            else:
                parsed = json.loads(raw_text)
            if self._is_valid_budget_payload(parsed):
                context_sensitive = any(
                    [timeline.strip(), project_type.strip(), service_type.strip(), project_goal.strip(), objective.strip(), notes.strip(), frameworks.strip()]
                )
                return self._merge_with_context_adjustment(
                    parsed,
                    dynamic_estimate,
                    context_sensitive=context_sensitive,
                    prefer_context_options=not staged_pricing
                )
            return dynamic_estimate
        except Exception as e:
            logger.error(f"Financial Analyzer Error: {e}")
            return dynamic_estimate

class LogoManager:
    @staticmethod
    def _create_fallback_logo(client_name: str) -> io.BytesIO:
        initials = "".join([w[0] for w in re.findall(r"[A-Za-z0-9]+", client_name)[:3]]).upper() or "CL"
        canvas = Image.new('RGB', (320, 320), color=(236, 242, 255))
        draw = ImageDraw.Draw(canvas)
        draw.rounded_rectangle((16, 16, 304, 304), radius=36, outline=(37, 99, 235), width=8)
        try:
            font = ImageFont.truetype("DejaVuSans-Bold.ttf", 110)
        except Exception:
            font = ImageFont.load_default()
        try:
            bbox = draw.textbbox((0, 0), initials, font=font)
            w = bbox[2] - bbox[0]
            h = bbox[3] - bbox[1]
        except Exception:
            w, h = draw.textsize(initials, font=font)
        draw.text(((320 - w) / 2, (320 - h) / 2 - 6), initials, fill=(15, 23, 42), font=font)
        out = io.BytesIO()
        canvas.save(out, format='PNG')
        out.seek(0)
        return out

    @staticmethod
    def get_logo_and_color(client_name: str) -> Tuple[Optional[io.BytesIO], Tuple[int, int, int]]:
        if not Researcher._has_serper_key():
            return LogoManager._create_fallback_logo(client_name), DEFAULT_COLOR
        try:
            url = "https://google.serper.dev/images"
            payload = json.dumps({"q": f"{client_name} corporate logo png transparent", "num": 3})
            headers = {'X-API-KEY': SERPER_API_KEY, 'Content-Type': 'application/json'}
            res = requests.post(url, headers=headers, data=payload, timeout=8).json()
            
            if 'images' in res and res['images']:
                for item in res['images']:
                    try:
                        img_resp = requests.get(item['imageUrl'], headers={'User-Agent': 'Mozilla/5.0'}, timeout=5)
                        if img_resp.status_code == 200:
                            stream = io.BytesIO(img_resp.content)
                            img = Image.open(stream)
                            
                            # Normalize to PNG so python-docx can embed it reliably.
                            if img.mode in ("RGBA", "LA", "P"):
                                normalized = Image.new("RGBA", img.size, (255, 255, 255, 0))
                                normalized.paste(img, (0, 0), img if img.mode in ("RGBA", "LA") else None)
                                img = normalized.convert('RGB')
                            else:
                                img = img.convert('RGB')

                            img.thumbnail((150, 150))
                            dom_color = list(map(int, ImageStat.Stat(img).mean[:3]))

                            luminance = 0.299 * dom_color[0] + 0.587 * dom_color[1] + 0.114 * dom_color[2]
                            if luminance > 120:
                                factor = 120 / luminance
                                dom_color = [max(0, min(255, int(c * factor))) for c in dom_color]

                            # Use normalized bytes, not source bytes, to avoid format errors.
                            png_stream = io.BytesIO()
                            img.save(png_stream, format='PNG')
                            png_stream.seek(0)
                            return png_stream, tuple(dom_color)
                    except Exception:
                        continue
        except Exception as e:
            logger.warning(f"Logo Retrieval Error: {e}")
        return LogoManager._create_fallback_logo(client_name), DEFAULT_COLOR


# Document rendering utilities.
class StyleEngine:
    @staticmethod
    def apply_document_styles(doc: Document) -> None:
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        pf = style.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = 1.15
        pf.space_after = Pt(8) 
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for section in doc.sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(2.54)
            section.right_margin = Cm(2.54)

class ChartEngine:
    @staticmethod
    def _to_matplotlib_rgb(theme_color: Tuple[int, int, int]) -> Tuple[float, float, float]:
        return tuple(c/255 for c in theme_color)

    @staticmethod
    def create_gantt_chart(data_str: str, theme_color: Tuple[int, int, int]) -> Optional[io.BytesIO]:
        try:
            parts = data_str.split('|')
            if len(parts) == 3:
                title_str, unit_str, raw_data = parts[0].strip(), parts[1].strip(), parts[2].strip()
            else:
                title_str, unit_str, raw_data = "Timeline", "Waktu", data_str
            tasks = []
            for p in raw_data.split(';'):
                t_parts = p.split(',')
                if len(t_parts) >= 3:
                    tasks.append({"task": t_parts[0].strip(), "start": float(re.sub(r'[^\d.]', '', t_parts[1])), "dur": float(re.sub(r'[^\d.]', '', t_parts[2]))})
            if not tasks: return None
            tasks = tasks[::-1] 
            
            fig, ax = plt.subplots(figsize=(8.5, max(4, len(tasks)*0.8)))
            for i, task in enumerate(tasks):
                rect = patches.FancyBboxPatch((task['start'], i-0.3), task['dur'], 0.6, boxstyle="round,pad=0.02", ec="#ffffff", fc=ChartEngine._to_matplotlib_rgb(theme_color), alpha=0.9, lw=1.5)
                ax.add_patch(rect)
                ax.text(task['start'] + (task['dur'] / 2), i, f"{task['dur']:g} {unit_str}", ha='center', va='center', color='white', fontweight='bold', fontsize=9)
            
            names = [t['task'] for t in tasks]
            ax.set_yticks(range(len(names)))
            ax.set_yticklabels(names, fontsize=10)
            ax.set_title(title_str, fontsize=13, fontweight='bold', pad=20)
            ax.grid(axis='x', linestyle='--', alpha=0.5)
            max_x = max([t['start'] + t['dur'] for t in tasks])
            ax.set_xlim(0, max_x + (max_x * 0.1))
            ax.set_ylim(-0.6, len(names)-0.4)
            
            import matplotlib.ticker as ticker
            ax.xaxis.set_major_locator(ticker.MaxNLocator(integer=True))
            
            img = io.BytesIO()
            plt.savefig(img, format='png', bbox_inches='tight', dpi=150)
            plt.close()
            img.seek(0)
            return img
        except Exception:
            return None

class DocumentBuilder:
    @staticmethod
    def _append_text_run(paragraph, text: str, bold: bool = False, italic: bool = False) -> None:
        cleaned = re.sub(r'\s+', ' ', text or '').strip()
        if not cleaned:
            return

        # Keep spacing between text runs stable so inline formatting stays readable.
        if paragraph.runs:
            last_text = paragraph.runs[-1].text or ""
            if last_text and not last_text.endswith((" ", "\t", "\n", "(", "[", "/")) and not cleaned.startswith((".", ",", ";", ":", ")", "]", "%")):
                cleaned = " " + cleaned

        run = paragraph.add_run(cleaned)
        run.bold = bold
        run.italic = italic

    @staticmethod
    def _style_num_id(doc: Document, style_name: str) -> Optional[int]:
        try:
            style = doc.styles[style_name]
        except KeyError:
            return None
        ppr = style.element.pPr
        if ppr is None or ppr.numPr is None or ppr.numPr.numId is None:
            return None
        try:
            return int(ppr.numPr.numId.val)
        except Exception:
            return None

    @staticmethod
    def _create_list_num_id(doc: Document, style_name: str) -> Optional[int]:
        style_num_id = DocumentBuilder._style_num_id(doc, style_name)
        if style_num_id is None:
            return None

        numbering = doc.part.numbering_part.numbering_definitions._numbering
        abstract_num_id = None
        for num in numbering.findall(qn('w:num')):
            raw_num_id = num.get(qn('w:numId'))
            try:
                current_num_id = int(raw_num_id) if raw_num_id is not None else None
            except (TypeError, ValueError):
                current_num_id = None
            if current_num_id != style_num_id:
                continue
            abstract = num.find(qn('w:abstractNumId'))
            if abstract is None:
                continue
            raw_abstract = abstract.get(qn('w:val'))
            try:
                abstract_num_id = int(raw_abstract) if raw_abstract is not None else None
            except (TypeError, ValueError):
                abstract_num_id = None
            break

        if abstract_num_id is None:
            return None

        existing_num_ids: List[int] = []
        for num in numbering.findall(qn('w:num')):
            raw_num_id = num.get(qn('w:numId'))
            try:
                if raw_num_id is not None:
                    existing_num_ids.append(int(raw_num_id))
            except (TypeError, ValueError):
                continue
        next_num_id = (max(existing_num_ids) + 1) if existing_num_ids else (style_num_id + 1)

        num = OxmlElement('w:num')
        num.set(qn('w:numId'), str(next_num_id))
        abstract_ref = OxmlElement('w:abstractNumId')
        abstract_ref.set(qn('w:val'), str(abstract_num_id))
        num.append(abstract_ref)
        numbering.append(num)
        return next_num_id

    @staticmethod
    def _apply_list_num_id(paragraph, num_id: int, level: int = 0) -> None:
        p = paragraph._p
        ppr = p.get_or_add_pPr()
        num_pr = ppr.find(qn('w:numPr'))
        if num_pr is None:
            num_pr = OxmlElement('w:numPr')
            ppr.append(num_pr)

        ilvl = num_pr.find(qn('w:ilvl'))
        if ilvl is None:
            ilvl = OxmlElement('w:ilvl')
            num_pr.append(ilvl)
        ilvl.set(qn('w:val'), str(level))

        num_id_el = num_pr.find(qn('w:numId'))
        if num_id_el is None:
            num_id_el = OxmlElement('w:numId')
            num_pr.append(num_id_el)
        num_id_el.set(qn('w:val'), str(num_id))

    @staticmethod
    def parse_html_to_docx(doc: Document, html_content: str, theme_color: Tuple[int, int, int]) -> None:
        soup = BeautifulSoup(html_content, 'html.parser')
        for element in soup.children:
            if element.name is None: continue
            if element.name in ['h1', 'h2', 'h3']:
                level = int(element.name[1])
                p = doc.add_heading(element.get_text().strip(), level=level)
                if level == 1:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.color.rgb = RGBColor(*theme_color)
                    run.font.name = 'Arial'
                    run.bold = True
            elif element.name == 'p':
                p = doc.add_paragraph()
                DocumentBuilder._process_inline_html(p, element)
            elif element.name in ['ul', 'ol']:
                # Use native Word list formatting so alignment and wrapping stay consistent.
                direct_items = element.find_all('li', recursive=False)
                style_name = "List Number" if element.name == 'ol' else "List Bullet"
                list_num_id = DocumentBuilder._create_list_num_id(doc, style_name)
                for idx, li in enumerate(direct_items, start=1):
                    # Avoid orphan markers like "3." with no text.
                    if not li.get_text(" ", strip=True):
                        continue
                    use_manual_fallback = False
                    try:
                        p = doc.add_paragraph(style=style_name)
                    except KeyError:
                        p = doc.add_paragraph()
                        use_manual_fallback = True

                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(4)
                    if list_num_id is not None:
                        DocumentBuilder._apply_list_num_id(p, list_num_id, level=0)
                    elif use_manual_fallback:
                        marker = f"{idx}. " if element.name == 'ol' else "• "
                        p.add_run(marker).bold = True
                    DocumentBuilder._process_inline_html(p, li)
            elif element.name == 'table':
                rows = element.find_all('tr')
                if not rows: continue
                max_cols = max([len(r.find_all(['td', 'th'])) for r in rows])
                table = doc.add_table(rows=len(rows), cols=max_cols)
                table.style = 'Table Grid'
                for i, row in enumerate(rows):
                    cols = row.find_all(['td', 'th'])
                    for j, col in enumerate(cols):
                        if j < max_cols:
                            cell = table.cell(i, j)
                            cell._element.clear_content()
                            p = cell.add_paragraph()
                            DocumentBuilder._process_inline_html(p, col)

    @staticmethod
    def _process_inline_html(paragraph, element):
        for child in element.children:
            if child.name in ['strong', 'b']:
                DocumentBuilder._append_text_run(paragraph, child.get_text(" ", strip=True), bold=True)
            elif child.name in ['em', 'i']:
                DocumentBuilder._append_text_run(paragraph, child.get_text(" ", strip=True), italic=True)
            elif child.name == 'br':
                paragraph.add_run("\n")
            elif child.name is None:
                DocumentBuilder._append_text_run(paragraph, str(child))
            else:
                DocumentBuilder._process_inline_html(paragraph, child)

    @staticmethod
    def process_content(doc: Document, raw_text: str, theme_color: Tuple[int, int, int], chapter_title: str) -> None:
        clean_lines = []
        in_table = False
        for line in raw_text.split('\n'):
            line = line.strip()
            if line.startswith('[[GANTT:') and line.endswith(']]'):
                data = line.replace('[[GANTT:', '').replace(']]', '').strip()
                img = ChartEngine.create_gantt_chart(data, theme_color)
                if img: doc.add_paragraph().add_run().add_picture(img, width=Inches(6))
                continue
            if line.startswith('|'):
                if not in_table and clean_lines and clean_lines[-1] != "":
                    clean_lines.append("")
                in_table = True
            else:
                in_table = False
            clean_lines.append(line)
            
        html = markdown.markdown("\n".join(clean_lines), extensions=['tables', 'sane_lists'])
        DocumentBuilder.parse_html_to_docx(doc, html, theme_color)


class ProposalGenerator:
    DEFAULT_CHAPTER_TARGET_WORDS = 700
    BASE_CHAPTER_FLOOR_WORDS = 220
    CLOSING_CHAPTER_FLOOR_WORDS = 170
    BASE_COMPRESSION_FLOOR_WORDS = 240
    CLOSING_COMPRESSION_FLOOR_WORDS = 180

    def __init__(self, kb_instance: KnowledgeBase) -> None:
        self.ollama = Client(host=OLLAMA_HOST)
        self.kb = kb_instance
        self.io_pool = concurrent.futures.ThreadPoolExecutor(max_workers=10)
        self.firm_api = FirmAPIClient()
        self._research_cache: Dict[str, Dict[str, str]] = {}
        self._proposal_contract_cache: Dict[str, str] = {}
        self._chapter_context_cache: Dict[str, Dict[str, str]] = {}

    @staticmethod
    def _target_words(chapter: Dict[str, Any]) -> int:
        m = re.search(r'Target:\s*(\d+)\s*words', chapter.get('length_intent', ''), re.IGNORECASE)
        return int(m.group(1)) if m else ProposalGenerator.DEFAULT_CHAPTER_TARGET_WORDS

    @staticmethod
    def _word_count(text: str) -> int:
        return len(re.findall(r'\b\w+\b', text))

    @staticmethod
    def _rewrite_length_intent(length_intent: str, target_words: int) -> str:
        if not length_intent:
            return f"Target: {target_words} words."
        if re.search(r'Target:\s*\d+\s*words', length_intent, re.IGNORECASE):
            return re.sub(
                r'Target:\s*\d+\s*words',
                f"Target: {int(target_words)} words",
                length_intent,
                flags=re.IGNORECASE
            )
        return f"{length_intent.rstrip('.')} (Target: {int(target_words)} words)."

    def _content_word_budget(self) -> int:
        usable_pages = max(1, MAX_PROPOSAL_PAGES - RESERVED_NON_CONTENT_PAGES - PAGE_SAFETY_BUFFER)
        return int(usable_pages * ESTIMATED_WORDS_PER_PAGE)

    @classmethod
    def _chapter_floor_words(cls, chapter_id: str, for_compression: bool = False) -> int:
        if chapter_id == "c_closing":
            return cls.CLOSING_COMPRESSION_FLOOR_WORDS if for_compression else cls.CLOSING_CHAPTER_FLOOR_WORDS
        return cls.BASE_COMPRESSION_FLOOR_WORDS if for_compression else cls.BASE_CHAPTER_FLOOR_WORDS

    def _chapter_word_targets(self, chapters: List[Dict[str, Any]]) -> Dict[str, int]:
        base_targets = {chapter["id"]: self._target_words(chapter) for chapter in chapters}
        total_base = sum(base_targets.values())
        budget = self._content_word_budget()
        if total_base <= 0:
            return {chapter["id"]: 500 for chapter in chapters}
        if total_base <= budget:
            return base_targets

        # Scale all chapters down proportionally while keeping minimum readable length.
        scaled: Dict[str, int] = {}
        for chapter in chapters:
            base = base_targets[chapter["id"]]
            floor = self._chapter_floor_words(chapter["id"], for_compression=False)
            scaled[chapter["id"]] = max(floor, int(base * budget / total_base))

        # If still over budget, cut from the longest chapters first.
        overflow = sum(scaled.values()) - budget
        if overflow > 0:
            ordered_ids = sorted(scaled.keys(), key=lambda cid: scaled[cid], reverse=True)
            for cid in ordered_ids:
                if overflow <= 0:
                    break
                floor = self._chapter_floor_words(cid, for_compression=False)
                reducible = max(0, scaled[cid] - floor)
                cut = min(reducible, overflow)
                scaled[cid] -= cut
                overflow -= cut
        return scaled

    def _estimated_pages(self, total_words: int) -> int:
        content_pages = max(1, (total_words + ESTIMATED_WORDS_PER_PAGE - 1) // ESTIMATED_WORDS_PER_PAGE)
        return RESERVED_NON_CONTENT_PAGES + content_pages

    @staticmethod
    def _cache_key(*parts: Any) -> str:
        return "||".join([str(p).strip().lower() for p in parts])

    @staticmethod
    def _cache_put(cache: Dict[str, Any], key: str, value: Any, max_size: int = 128) -> None:
        if key in cache:
            cache.pop(key, None)
        cache[key] = value
        while len(cache) > max_size:
            cache.pop(next(iter(cache)))

    @staticmethod
    def _has_external_evidence(osint_block: str) -> bool:
        return any(line.strip().startswith("Sumber eksternal") for line in (osint_block or "").splitlines())

    @staticmethod
    def _extract_osint_facts(osint_block: str, max_items: int = 3) -> List[Dict[str, str]]:
        facts: List[Dict[str, str]] = []
        pattern = (
            r"fakta=(?P<fact>.*?)\s*\|\s*sumber=(?P<title>.*?)\s*\|\s*url=(?P<url>.*?)\s*"
            r"\|\s*sitasi_apa=(?P<citation>\([^)]+\))$"
        )
        for raw_line in (osint_block or "").splitlines():
            line = raw_line.strip()
            if not line.startswith("Sumber eksternal"):
                continue
            match = re.search(pattern, line)
            if not match:
                continue
            fact = re.sub(r"\s+", " ", (match.group("fact") or "")).strip()
            title = re.sub(r"\s+", " ", (match.group("title") or "")).strip()
            url = (match.group("url") or "").strip()
            citation = (match.group("citation") or "").strip()
            if not fact:
                continue
            facts.append({
                "fact": fact[:240],
                "title": title,
                "url": url,
                "citation": citation,
            })
            if len(facts) >= max_items:
                break
        return facts

    @staticmethod
    def _infer_industry(client: str, project: str, notes: str, regulations: str) -> str:
        combined = " ".join([client or "", project or "", notes or "", regulations or ""]).lower()
        rules = [
            ("Perbankan", ["bank", "bri", "bca", "mandiri", "bni", "btn", "fintech", "kredit"]),
            ("Telekomunikasi", ["telkom", "telkomsel", "indosat", "xl", "axiata", "operator"]),
            ("Energi & Utilitas", ["pertamina", "pln", "energi", "listrik", "oil", "gas"]),
            ("Ritel & E-Commerce", ["tokopedia", "e-commerce", "retail", "marketplace", "goto"]),
            ("Transportasi & Aviasi", ["garuda", "aviation", "airline", "logistik", "transport"]),
            ("Pemerintah & BUMN", ["kementerian", "dinas", "pemprov", "pemkab", "bumn", "spbe"]),
            ("Manufaktur & Tambang", ["manufaktur", "adaro", "antam", "vale", "smelter", "mining"]),
        ]
        for label, tokens in rules:
            if any(token in combined for token in tokens):
                return label
        return "Lintas Industri"

    @staticmethod
    def _industry_terms(industry: str) -> List[str]:
        terms_map = {
            "Perbankan": ["nasabah", "core banking", "risk appetite", "kepatuhan POJK", "fraud control"],
            "Telekomunikasi": ["subscriber", "network availability", "service assurance", "NOC", "customer churn"],
            "Energi & Utilitas": ["operational reliability", "asset integrity", "outage window", "HSSE", "dispatch"],
            "Ritel & E-Commerce": ["conversion", "basket size", "checkout funnel", "retention cohort", "omnichannel"],
            "Transportasi & Aviasi": ["on-time performance", "flight disruption", "ground operation", "passenger experience"],
            "Pemerintah & BUMN": ["akuntabilitas", "SPBE", "tata kelola", "layanan publik", "audit trail"],
            "Manufaktur & Tambang": ["supply continuity", "production throughput", "downtime pabrik", "safety compliance"],
        }
        return terms_map.get(industry, ["operational excellence", "risk control", "business value", "governance"])

    @staticmethod
    def _extract_metric_keywords(kpi_blueprint: List[str]) -> List[str]:
        stopwords = {"dan", "yang", "untuk", "dengan", "dari", "pada", "target", "baseline", "dalam", "bulan"}
        keywords: List[str] = []
        seen: Set[str] = set()
        for line in kpi_blueprint or []:
            for token in re.findall(r"[A-Za-z]{3,}", (line or "").lower()):
                if token in stopwords or token in seen:
                    continue
                seen.add(token)
                keywords.append(token)
                if len(keywords) >= 10:
                    return keywords
        return keywords

    @staticmethod
    def _extract_anchor_keywords(facts: List[Dict[str, str]], max_terms: int = 10) -> List[str]:
        stopwords = {
            "yang", "dengan", "untuk", "dari", "pada", "tahun", "bulan", "sebagai", "dan",
            "atau", "oleh", "this", "that", "from", "with", "into", "their"
        }
        seen: Set[str] = set()
        terms: List[str] = []
        for fact in facts or []:
            for token in re.findall(r"[A-Za-z]{4,}", (fact.get("fact", "") or "").lower()):
                if token in stopwords or token in seen:
                    continue
                seen.add(token)
                terms.append(token)
                if len(terms) >= max_terms:
                    return terms
        return terms

    @classmethod
    def _build_kpi_blueprint(
        cls,
        project_goal: str,
        notes: str,
        timeline: str,
        industry: str,
        client: str
    ) -> List[str]:
        combined = " ".join([project_goal or "", notes or ""]).lower()
        months = FinancialAnalyzer._duration_to_months(timeline)
        horizon = f"{months:.0f} bulan" if months else (timeline or "periode proyek")

        kpis: List[str] = []
        if any(token in combined for token in ["mau", "monthly active", "active users", "adopsi"]):
            kpis.append(f"MAU/Adopsi kanal digital: target pertumbuhan terukur dalam {horizon}.")
        if any(token in combined for token in ["nps", "kepuasan", "customer experience", "csat"]):
            kpis.append(f"Customer sentiment (NPS/CSAT): target kenaikan konsisten dalam {horizon}.")
        if any(token in combined for token in ["churn", "retensi", "retention"]):
            kpis.append(f"Retention/Churn: turunkan churn dan naikkan retensi pada horizon {horizon}.")
        if any(token in combined for token in ["downtime", "incident", "availability", "sla", "mttr"]):
            kpis.append(f"Reliability operasi (Availability/MTTR/Incident): perbaikan stabil dalam {horizon}.")
        if any(token in combined for token in ["biaya", "cost", "efisiensi", "opex", "capex"]):
            kpis.append(f"Efisiensi biaya delivery: kontrol biaya per fase dan deviasi anggaran dalam {horizon}.")

        if not kpis:
            default_by_industry = {
                "Perbankan": [
                    f"Pertumbuhan transaksi digital nasabah dalam {horizon}.",
                    f"Kepatuhan & risiko operasional (audit findings) membaik dalam {horizon}.",
                    f"Ketersediaan layanan digital meningkat dan insiden prioritas menurun dalam {horizon}.",
                ],
                "Telekomunikasi": [
                    f"Service availability dan pengalaman pelanggan meningkat dalam {horizon}.",
                    f"Churn pelanggan menurun dan quality of service membaik dalam {horizon}.",
                    f"Efisiensi operasi network support meningkat dalam {horizon}.",
                ],
            }
            kpis.extend(default_by_industry.get(industry, [
                f"KPI outcome bisnis utama {client} bergerak positif dalam {horizon}.",
                f"Kualitas operasi (incident, SLA, lead time) membaik dalam {horizon}.",
                f"Efektivitas eksekusi proyek (on-time, on-budget, on-scope) tercapai dalam {horizon}.",
            ]))

        # Keep concise but actionable.
        return kpis[:4]

    @classmethod
    def _build_personalization_pack(
        cls,
        client: str,
        project: str,
        project_goal: str,
        timeline: str,
        notes: str,
        regulations: str,
        research_bundle: Dict[str, str],
        relationship_context: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        industry = cls._infer_industry(client, project, notes, regulations)
        track_record = research_bundle.get("track_record", "")
        collaboration = research_bundle.get("collaboration", "")
        news = research_bundle.get("news", "")

        relationship_context = relationship_context or {}
        relationship_source = str(relationship_context.get("source") or "").strip().lower()
        relationship_mode = str(relationship_context.get("mode") or "").strip().lower()
        if relationship_mode not in {"existing", "new"}:
            relationship_has_evidence = cls._has_external_evidence(collaboration)
            relationship_mode = "existing" if relationship_has_evidence else "new"
            relationship_source = "osint" if relationship_has_evidence else relationship_source
        if relationship_source == "internal_api":
            relationship_guidance = (
                "Gunakan narasi continuity partnership hanya jika data internal memang menunjukkan hubungan kerja sebelumnya."
                if relationship_mode == "existing"
                else "Gunakan narasi kemitraan baru. Jangan klaim hubungan sebelumnya tanpa data internal."
            )
        else:
            relationship_guidance = (
                "Gunakan narasi continuity partnership berbasis bukti publik kolaborasi yang ada."
                if relationship_mode == "existing"
                else "Gunakan narasi kemitraan baru. Jangan klaim pernah bekerja sama tanpa bukti publik."
            )

        news_facts = cls._extract_osint_facts(news, max_items=2)
        track_facts = cls._extract_osint_facts(track_record, max_items=2)
        initiative_facts: List[Dict[str, str]] = []
        seen_citations: Set[str] = set()
        for item in news_facts + track_facts:
            citation = item.get("citation", "")
            if citation in seen_citations:
                continue
            seen_citations.add(citation)
            initiative_facts.append(item)
            if len(initiative_facts) >= 3:
                break

        kpi_blueprint = cls._build_kpi_blueprint(
            project_goal=project_goal,
            notes=notes,
            timeline=timeline,
            industry=industry,
            client=client
        )
        terminology = cls._industry_terms(industry)

        profile_summary = (
            f"Industri klien terdeteksi: {industry}. "
            f"Fokus inisiatif: {project}. "
            f"Model hubungan: {'kelanjutan kolaborasi' if relationship_mode == 'existing' else 'inisiasi kemitraan baru'}."
        )

        return {
            "industry": industry,
            "relationship_mode": relationship_mode,
            "relationship_source": relationship_source or "osint",
            "relationship_guidance": relationship_guidance,
            "kpi_blueprint": kpi_blueprint,
            "kpi_keywords": cls._extract_metric_keywords(kpi_blueprint),
            "terminology": terminology[:6],
            "initiative_facts": initiative_facts,
            "anchor_citations": [item.get("citation", "") for item in initiative_facts if item.get("citation")],
            "anchor_keywords": cls._extract_anchor_keywords(initiative_facts),
            "profile_summary": profile_summary,
        }

    @staticmethod
    def _normalize_external_citation(domain: str, year: str) -> str:
        domain_clean = (domain or "").strip().lower()
        year_clean = (year or "").strip().lower()
        if year_clean in {"nd", "n.d", "n.d"}:
            year_clean = "n.d."
        return f"({domain_clean}, {year_clean})"

    @classmethod
    def _extract_external_citations(cls, text: str) -> List[str]:
        pattern = r"\(([A-Za-z0-9.-]+\.[A-Za-z]{2,}),\s*(\d{4}|n\.d\.)\)"
        citations = []
        for domain, year in re.findall(pattern, text or "", flags=re.IGNORECASE):
            citations.append(cls._normalize_external_citation(domain, year))
        return citations

    @classmethod
    def _collect_allowed_external_citations(cls, research_bundle: Dict[str, str]) -> Set[str]:
        if not research_bundle:
            return set()
        merged = "\n".join([str(value) for value in research_bundle.values()])
        return set(cls._extract_external_citations(merged))

    @classmethod
    def _clean_external_citations(cls, content: str, allowed_external_citations: Set[str]) -> str:
        allowed = set(allowed_external_citations or set())
        pattern = r"\(([A-Za-z0-9.-]+\.[A-Za-z]{2,}),\s*(\d{4}|n\.d\.)\)"

        def replace_invalid(match: re.Match) -> str:
            citation = cls._normalize_external_citation(match.group(1), match.group(2))
            return match.group(0) if citation in allowed else ""

        cleaned = re.sub(pattern, replace_invalid, content or "", flags=re.IGNORECASE)
        cleaned = re.sub(r"[ \t]{2,}", " ", cleaned)
        cleaned = re.sub(r"\(\s*\)", "", cleaned)
        cleaned = re.sub(r"\s+([,.;:])", r"\1", cleaned)
        return cleaned

    @staticmethod
    def _verified_firm_contact_block(firm_profile: Optional[Dict[str, Any]]) -> str:
        contact_lines = FirmAPIClient.build_contact_lines(firm_profile)
        if not contact_lines:
            return ""
        return "\n".join([f"- {line}" for line in contact_lines])

    @classmethod
    def _inject_verified_firm_contact(
        cls,
        content: str,
        firm_profile: Optional[Dict[str, Any]]
    ) -> str:
        block = cls._verified_firm_contact_block(firm_profile)
        if not block:
            return content

        existing_lines = FirmAPIClient.build_contact_lines(firm_profile)
        if all(line in (content or "") for line in existing_lines):
            return content

        contact_section = f"### Kontak Resmi Terverifikasi\n{block}"
        heading_pattern = r"(?im)^\s*##\s*Informasi Kontak dan Langkah Lanjutan\s*$"
        match = re.search(heading_pattern, content or "")
        if match:
            insert_at = match.end()
            return (content or "")[:insert_at] + "\n\n" + contact_section + (content or "")[insert_at:]

        trimmed = (content or "").rstrip()
        if trimmed:
            trimmed += "\n\n"
        return (
            f"{trimmed}## Informasi Kontak dan Langkah Lanjutan\n\n"
            f"{contact_section}"
        )

    def _resolve_chapters(self, chapter_id: Optional[str]) -> List[Dict[str, Any]]:
        normalized_id = (chapter_id or "").strip()
        if not normalized_id or normalized_id.lower() in {"all", "semua"}:
            return UNIVERSAL_STRUCTURE

        selected = [chapter for chapter in UNIVERSAL_STRUCTURE if chapter["id"] == normalized_id]
        if selected:
            return selected

        normalized = normalized_id.lower()
        selected = [chapter for chapter in UNIVERSAL_STRUCTURE if chapter["title"].strip().lower() == normalized]
        if selected:
            return selected

        raise ValueError(f"Unknown chapter_id: {normalized_id}")

    def build_preview_outline(self, data: Dict[str, Any]) -> List[Dict[str, Any]]:
        chapter_id = (data or {}).get("chapter_id")
        try:
            chapters = self._resolve_chapters(chapter_id)
        except ValueError:
            chapters = UNIVERSAL_STRUCTURE

        client = (data or {}).get("nama_perusahaan", "Klien")
        objective = (data or {}).get("konteks_organisasi", "").strip() or "tujuan proyek belum diisi"
        issues = (data or {}).get("permasalahan", "").strip() or "pain points belum diisi"
        need_type = (data or {}).get("klasifikasi_kebutuhan", "").strip() or "belum dipilih"
        project_type = (data or {}).get("jenis_proyek", "").strip() or "belum dipilih"
        service_type = (data or {}).get("jenis_proposal", "").strip() or "belum dipilih"
        frameworks = (data or {}).get("potensi_framework", "").strip() or "belum dipilih"
        timeline = (data or {}).get("estimasi_waktu", "").strip() or "belum ditentukan"
        budget = (data or {}).get("estimasi_biaya", "").strip() or "belum ditentukan"

        preview_map = {
            "c_1": f"Menetapkan konteks organisasi {client} dan objektif inisiatif: {objective}.",
            "c_2": f"Mengurai kebutuhan dan akar masalah klien berdasarkan pain points: {issues}.",
            "c_3": f"Mengkategorikan kebutuhan ke {need_type} lalu memvalidasi jenis proyek {project_type}.",
            "c_4": f"Menautkan kebutuhan klien dengan framework/regulasi utama: {frameworks}.",
            "c_5": f"Menjelaskan pemilihan metodologi delivery untuk layanan {service_type}.",
            "c_6": "Mendetailkan target state, output, dan deliverable solusi yang dapat dieksekusi.",
            "c_7": f"Menyusun rencana fase, milestone, dan deliverable berdasarkan durasi {timeline}.",
            "c_8": "Merumuskan governance proyek: forum keputusan, eskalasi isu, dan quality gate.",
            "c_9": f"Menetapkan struktur tim dan kapabilitas yang dibutuhkan untuk model {service_type}.",
            "c_10": f"Mendefinisikan model pembiayaan, termin pembayaran, dan batasan scope dengan estimasi {budget}.",
            "c_closing": f"Menutup proposal dengan apresiasi kemitraan, kontak resmi {WRITER_FIRM_NAME}, dan langkah tindak lanjut bersama {client}.",
        }

        return [
            {
                "id": chapter["id"],
                "title": chapter["title"],
                "preview": preview_map.get(chapter["id"], "Ringkasan konten bab akan disesuaikan dengan konteks klien."),
                "subsections": chapter["subs"],
            }
            for chapter in chapters
        ]

    def _get_research_bundle(self, base_client: str, regulations: str, include_collaboration: bool = True) -> Dict[str, str]:
        key = self._cache_key("research", base_client, regulations, str(include_collaboration))
        cached = self._research_cache.get(key)
        if cached:
            return cached

        futures = {
            "profile": self.io_pool.submit(Researcher.get_entity_profile, base_client),
            "news": self.io_pool.submit(Researcher.get_latest_client_news, base_client),
            "track_record": self.io_pool.submit(Researcher.get_client_track_record, base_client),
            "regulations": self.io_pool.submit(Researcher.get_regulatory_data, regulations)
        }
        if include_collaboration:
            futures["collaboration"] = self.io_pool.submit(
                Researcher.get_client_writer_collaboration, base_client, WRITER_FIRM_NAME
            )
        try:
            bundle = {
                "profile": futures["profile"].result(timeout=8),
                "news": futures["news"].result(timeout=8),
                "track_record": futures["track_record"].result(timeout=8),
                "collaboration": futures["collaboration"].result(timeout=8) if include_collaboration else "",
                "regulations": futures["regulations"].result(timeout=8),
            }
        except Exception:
            bundle = {
                "profile": f"Data profil terbaru {base_client} terbatas (sumber daring, n.d.).",
                "news": f"Data berita terbaru {base_client} terbatas (sumber daring, n.d.).",
                "track_record": f"Data track record {base_client} terbatas (sumber daring, n.d.).",
                "collaboration": (
                    f"Data histori kolaborasi {WRITER_FIRM_NAME} dengan {base_client} terbatas (sumber daring, n.d.)."
                    if include_collaboration else ""
                ),
                "regulations": "Data regulasi terbatas (sumber daring, n.d.)."
            }

        self._cache_put(self._research_cache, key, bundle, max_size=96)
        return bundle

    def _build_proposal_contract(
        self,
        client: str,
        project: str,
        budget: str,
        service_type: str,
        project_goal: str,
        project_type: str,
        timeline: str,
        notes: str,
        regulations: str,
        selected_chapters: List[Dict[str, Any]],
        research_bundle: Dict[str, str],
        firm_data: Dict[str, str],
        personalization_pack: Dict[str, Any]
    ) -> str:
        cache_key = self._cache_key(
            "contract", client, project, budget, service_type, project_goal,
            project_type, timeline, notes, regulations, "|".join([c["id"] for c in selected_chapters]),
            personalization_pack.get("industry", ""),
            personalization_pack.get("relationship_mode", ""),
            "|".join(personalization_pack.get("kpi_blueprint", []) or [])
        )
        cached = self._proposal_contract_cache.get(cache_key)
        if cached:
            return cached

        chapter_titles = ", ".join([c["title"] for c in selected_chapters])
        prompt = f"""
        Buat "Proposal Contract" ringkas untuk menjaga kualitas dan koherensi lintas bab.
        Konteks:
        - Klien: {client}
        - Inisiatif: {project}
        - Service Type: {service_type}
        - Jenis Proyek: {project_type}
        - Kebutuhan: {project_goal}
        - Durasi: {timeline}
        - Estimasi Biaya: {budget}
        - Pain Points: {notes}
        - Framework: {regulations}
        - Bab yang ditulis: {chapter_titles}
        - Baseline Metodologi: {firm_data.get('methodology', '')}
        - Baseline Team: {firm_data.get('team', '')}
        - Baseline Commercial: {firm_data.get('commercial', '')}
        - Profil OSINT: {research_bundle.get('profile', '')}
        - Berita OSINT: {research_bundle.get('news', '')}
        - Track Record OSINT: {research_bundle.get('track_record', '')}
        - Histori Kolaborasi {WRITER_FIRM_NAME} & {client}: {research_bundle.get('collaboration', '')}
        - Industry Pack: {personalization_pack.get('industry', 'Lintas Industri')}
        - Relationship Mode: {personalization_pack.get('relationship_mode', 'new')}
        - KPI Blueprint: {', '.join(personalization_pack.get('kpi_blueprint', []))}
        - Terminologi Prioritas: {', '.join(personalization_pack.get('terminology', []))}

        OUTPUT WAJIB (tanpa markdown code block, <= 220 kata):
        1) Narasi Inti (1-2 kalimat)
        2) Terminologi Kanonis (maks 6 istilah)
        3) Prinsip Konsistensi Antarbab (maks 5 butir)
        4) Larangan Gaya Tulis (maks 3 butir)
        """
        try:
            res = self.ollama.chat(
                model=LLM_MODEL,
                messages=[{'role': 'user', 'content': prompt}],
                options={'num_ctx': 16384, 'num_predict': 700, 'temperature': 0.15}
            )
            contract = (res.get('message', {}).get('content', '') or '').strip()
        except Exception:
            contract = ""

        if not contract:
            contract = (
                "Narasi Inti: Proposal harus menjawab kebutuhan bisnis klien secara konkret, terukur, dan eksekutabel.\n"
                "Terminologi Kanonis: deliverable, milestone, target state, governance, quality gate, risiko.\n"
                "Prinsip Konsistensi Antarbab: istilah konsisten, alur masalah-ke-solusi jelas, "
                "timeline sinkron dengan deliverable, tata kelola tegas, hindari repetisi.\n"
                "Larangan Gaya Tulis: filler generik, klaim tanpa dasar, paragraf tanpa tindakan."
            )

        self._cache_put(self._proposal_contract_cache, cache_key, contract, max_size=96)
        return contract

    def _build_chapter_prompt(
        self,
        chapter: Dict[str, Any],
        client: str,
        project: str,
        budget: str,
        service_type: str,
        project_goal: str,
        project_type: str,
        timeline: str,
        notes: str,
        regulations: str,
        firm_data: Dict[str, str],
        firm_profile: Dict[str, str],
        research_bundle: Dict[str, str],
        personalization_pack: Dict[str, Any],
        proposal_contract: str,
        target_words: int
    ) -> Dict[str, Any]:
        try:
            current_year = datetime.now().year
            profile_data = research_bundle.get('profile', '')
            track_record_data = research_bundle.get('track_record', '')
            collaboration_data = research_bundle.get('collaboration', '')
            global_data = "\n".join([item for item in [profile_data, track_record_data, collaboration_data] if item])
            client_news = research_bundle.get('news', '')
            regulation_data = research_bundle.get('regulations', '')
            allowed_external_citations = self._collect_allowed_external_citations(research_bundle)

            ctx_key = self._cache_key("chapter_ctx", client, project, budget, chapter.get('id', ''), chapter.get('keywords', ''))
            cached_ctx = self._chapter_context_cache.get(ctx_key)
            if cached_ctx:
                structured_row_data = cached_ctx.get("structured_row_data", "")
                rag_data = cached_ctx.get("rag_data", "")
            else:
                structured_row_data = self.kb.get_exact_context(client, project, budget)
                rag_data = self.kb.query(client, project, chapter['keywords'])
                self._cache_put(
                    self._chapter_context_cache,
                    ctx_key,
                    {"structured_row_data": structured_row_data, "rag_data": rag_data},
                    max_size=256
                )
            
            persona = PERSONAS.get(chapter.get('id', 'default'), PERSONAS['default'])
            subs = "\n".join([f"- {s}" for s in chapter['subs']])
            
            visual_prompt = ""
            if chapter.get('visual_intent') == "gantt":
                visual_prompt = f"Mandatory Timeline Visual: [[GANTT: Jadwal Pelaksanaan | Bulan | Fase 1,0,2; Fase 2,2,4]]. Total timeline: {timeline}."
            elif chapter.get('visual_intent') == "flowchart":
                visual_prompt = "Tambahkan alur tahapan metodologi dalam bentuk bullet bertingkat yang jelas (fase -> aktivitas -> output)."

            terminology_list = personalization_pack.get("terminology", []) or []
            kpi_blueprint = personalization_pack.get("kpi_blueprint", []) or []
            relationship_guidance = personalization_pack.get("relationship_guidance", "")
            relationship_mode = personalization_pack.get("relationship_mode", "new")
            relationship_source = personalization_pack.get("relationship_source", "osint")
            profile_summary = personalization_pack.get("profile_summary", "")
            initiative_facts = personalization_pack.get("initiative_facts", []) or []
            if initiative_facts:
                anchors_text = " | ".join([
                    f"{item.get('fact', '')} {item.get('citation', '')}".strip()
                    for item in initiative_facts[:3]
                ])
            else:
                anchors_text = "Data inisiatif publik terbatas; jangan membuat klaim inisiatif tanpa bukti."

            extra = (
                f"[PROPOSAL CONTRACT]\n{proposal_contract}\n"
                f"[GLOBAL] Proposal ini wajib mempertahankan kedalaman konten tingkat eksekutif dan total dokumen maksimal {MAX_PROPOSAL_PAGES} halaman. "
                "Setiap bab harus memiliki konteks spesifik klien, poin yang dapat ditindaklanjuti, dan tidak generik. Gunakan kombinasi numbering dan bullet yang rapi di setiap H2, namun tetap padat dan tidak banyak whitespace."
            )
            extra += f" [CLIENT_PROFILE_PACK] {profile_summary}"
            extra += f" [RELATIONSHIP_MODE] {relationship_mode}. {relationship_guidance}"
            extra += (
                f" [KPI_TAILORING] Gunakan KPI blueprint berikut sebagai baseline tailoring: "
                f"{' | '.join(kpi_blueprint) if kpi_blueprint else 'KPI belum spesifik, gunakan KPI operasional dan outcome bisnis yang terukur.'}"
            )
            extra += (
                f" [TERMINOLOGI_ADAPTASI] Gunakan istilah domain klien berikut secara natural: "
                f"{', '.join(terminology_list) if terminology_list else 'operational excellence, governance, risk control'}."
            )
            extra += (
                f" [INITIATIVE_ANCHORS] Wajib menyisipkan minimal satu anchor inisiatif klien dari daftar berikut: {anchors_text}"
            )
            extra += f" [OSINT_TRACK_RECORD] {track_record_data}"
            extra += (
                f" [{'INTERNAL_RELATIONSHIP' if relationship_source == 'internal_api' else 'OSINT_KOLABORASI'}] {collaboration_data} "
                + (
                    "Jangan pernah mengklaim pernah bekerja sama sebelumnya jika tidak ada data internal yang jelas."
                    if relationship_source == "internal_api"
                    else "Jangan pernah mengklaim pernah bekerja sama sebelumnya jika tidak ada bukti publik yang jelas."
                )
            )
            if chapter['id'] == 'c_1':
                extra += f" [FOCUS] Fokus pada latar belakang organisasi '{client}' dan tujuan proyek: '{project}'. Soroti driver bisnis utama: [{project_goal}]."
            elif chapter['id'] == 'c_2':
                extra += f" [FOCUS] Jabarkan kebutuhan/keinginan klien berdasarkan pain points berikut: '{notes}'. Gunakan analisis masalah yang tajam dan ringkas."
            elif chapter['id'] == 'c_3':
                extra += f" [FOCUS] Klasifikasikan kebutuhan ke Problem/Opportunity/Directive berdasarkan input: '{project_goal}'. Tetapkan jenis proyek: '{project_type}'."
            elif chapter['id'] == 'c_4':
                extra += f" [FOCUS] Gunakan framework/regulasi terpilih berikut sebagai acuan utama: '{regulations}'. Petakan langsung ke kebutuhan klien."
            elif chapter['id'] == 'c_5':
                extra += f" [FOCUS] Jelaskan alasan pemilihan metodologi untuk engagement '{service_type}' dan gunakan baseline metodologi internal: {firm_data['methodology']}."
            elif chapter['id'] == 'c_6':
                extra += f" [FOCUS] Turunkan metodologi menjadi solution design yang konkret: output, deliverable, dan target state yang dapat dieksekusi."
            elif chapter['id'] == 'c_7':
                extra += f" [FOCUS] Timeline harus sinkron dengan durasi proyek: '{timeline}'. Tampilkan aktivitas per fase, milestone, dan deliverable yang terukur."
            elif chapter['id'] == 'c_8':
                extra += " [FOCUS] Definisikan model tata kelola proyek: forum keputusan, frekuensi rapat, eskalasi isu, quality gate, dan kontrol progres."
            elif chapter['id'] == 'c_9':
                extra += f" [FOCUS] Uraikan struktur tim proyek untuk model layanan '{service_type}' dengan kapabilitas kunci, pengalaman, dan sertifikasi relevan. Referensi komposisi inti: {firm_data['team']}."
            elif chapter['id'] == 'c_10':
                extra += f" [FOCUS] Wajib menyajikan model pembiayaan dengan angka estimasi: {budget}. Sertakan termin pembayaran, model kerja, asumsi, eksklusi, dan terms komersial: {firm_data['commercial']}. Gunakan tabel markdown."
            elif chapter['id'] == 'c_closing':
                verified_contact_block = self._verified_firm_contact_block(firm_profile)
                extra += (
                    f" [FOCUS] Ini adalah bab penutup proposal. Jangan pernah menulis label 'BAB XI' atau variasinya. "
                    f"Tunjukkan apresiasi profesional kepada klien '{client}', tegaskan komitmen kolaborasi jangka panjang, "
                    f"dan berikan langkah tindak lanjut yang jelas dan actionable. "
                    f"Gunakan tone hangat, profesional, dan meyakinkan."
                )
                if verified_contact_block:
                    extra += (
                        " [CONTACT_POLICY] Cantumkan hanya detail kontak firma yang sudah terverifikasi berikut ini "
                        "secara persis, tanpa menambah atau memodifikasi data:\n"
                        f"{verified_contact_block}"
                    )
                else:
                    extra += (
                        " [CONTACT_POLICY] Selain nama firma, tidak ada detail kontak writer firm yang terverifikasi. "
                        "Jangan menulis alamat kantor, email, telepon, atau website yang tidak tersedia."
                    )

            if allowed_external_citations:
                allowed_list = ", ".join(sorted(allowed_external_citations))
                extra += (
                    f" [CITATION] Sitasi eksternal hanya boleh memakai daftar ini: {allowed_list}. "
                    f"Untuk klaim dari data internal, gunakan (Data Internal, {current_year}). "
                    "Dilarang membuat domain/sitasi eksternal baru di luar daftar dan dilarang memakai placeholder sitasi."
                )
            else:
                extra += (
                    f" [CITATION] Tidak ada sumber eksternal tervalidasi untuk bab ini. "
                    f"Dilarang menulis sitasi domain eksternal apa pun. Gunakan (Data Internal, {current_year}) "
                    "untuk klaim yang berasal dari data internal."
                )

            internal_citation_note = f"Gunakan sitasi internal: (Data Internal, {current_year})."
            structured_row_data_with_note = (
                f"{structured_row_data}\n{internal_citation_note}" if structured_row_data else internal_citation_note
            )
            rag_data_with_note = f"{rag_data}\n{internal_citation_note}" if rag_data else internal_citation_note

            prompt = PROPOSAL_SYSTEM_PROMPT.format(
                client=client, 
                writer_firm=WRITER_FIRM_NAME, 
                persona=persona,
                global_data=global_data, 
                client_news=client_news, 
                regulation_data=regulation_data,
                structured_row_data=structured_row_data_with_note,
                rag_data=rag_data_with_note,
                current_year=current_year,
                visual_prompt=visual_prompt, 
                extra_instructions=extra,
                chapter_title=chapter['title'], 
                sub_chapters=subs, 
                length_intent=self._rewrite_length_intent(chapter.get('length_intent', ''), target_words)
            )
            return {"prompt": prompt, "success": True}
        except Exception as e:
            return {"prompt": "", "success": False, "error": str(e)}

    @staticmethod
    def _contains_client_reference(content: str, client: str) -> bool:
        tokens = [
            t for t in re.findall(r"[A-Za-z0-9]{3,}", client)
            if t.lower() not in {"pt", "cv", "tbk"}
        ]
        if not tokens:
            return True
        return any(re.search(rf"\b{re.escape(tok)}\b", content, re.IGNORECASE) for tok in tokens[:3])

    def _evaluate_chapter_quality(
        self,
        chapter: Dict[str, Any],
        content: str,
        client: str,
        target_words: Optional[int] = None,
        allowed_external_citations: Optional[Set[str]] = None,
        personalization_pack: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        target_words = int(target_words or self._target_words(chapter))
        floor = max(140, int(self._chapter_floor_words(chapter.get("id", ""), for_compression=False) * 0.8))
        min_words = max(floor, int(target_words * 0.72))
        max_words = max(min_words + 90, int(target_words * 1.25))
        word_count = self._word_count(content)
        missing_h2 = [
            sub for sub in chapter.get('subs', [])
            if not re.search(rf"(?im)^\s*##\s*{re.escape(sub)}\s*$", content)
        ]
        has_numbered_list = bool(re.search(r"(?m)^\s*\d+\.\s+\S+", content))
        has_bullet_list = bool(re.search(r"(?m)^\s*[-*]\s+\S+", content))
        has_client_ref = self._contains_client_reference(content, client)
        has_required_visual = True
        if chapter.get('visual_intent') == "gantt":
            has_required_visual = "[[GANTT:" in content

        issues = []
        if missing_h2:
            issues.append("missing_h2")
        if word_count < min_words:
            issues.append("too_short")
        if word_count > max_words:
            issues.append("too_long")
        if not has_numbered_list or not has_bullet_list:
            issues.append("list_structure")
        if not has_required_visual:
            issues.append("missing_visual")
        if not has_client_ref:
            issues.append("missing_client_ref")

        invalid_external_citations: List[str] = []
        if allowed_external_citations is not None:
            cited_external = sorted(set(self._extract_external_citations(content)))
            allowed = set(allowed_external_citations)
            if not allowed and cited_external:
                invalid_external_citations = cited_external
            elif allowed:
                invalid_external_citations = [citation for citation in cited_external if citation not in allowed]
            if invalid_external_citations:
                issues.append("citation_policy")

        missing_personalization_signals: List[str] = []
        if personalization_pack:
            terminology = personalization_pack.get("terminology", []) or []
            term_hit = any(
                re.search(rf"\b{re.escape(term)}\b", content, re.IGNORECASE)
                for term in terminology if term
            ) if terminology else True
            if not term_hit:
                missing_personalization_signals.append("terminology")

            kpi_keywords = personalization_pack.get("kpi_keywords", []) or []
            kpi_hit = any(
                re.search(rf"\b{re.escape(token)}\b", content, re.IGNORECASE)
                for token in kpi_keywords if token
            ) if kpi_keywords else True
            if not kpi_hit:
                missing_personalization_signals.append("kpi")

            anchor_citations = personalization_pack.get("anchor_citations", []) or []
            anchor_hit = any(citation in content for citation in anchor_citations) if anchor_citations else True
            if not anchor_hit:
                missing_personalization_signals.append("initiative_anchor")

            if missing_personalization_signals:
                issues.append("missing_personalization")

        return {
            "issues": issues,
            "word_count": word_count,
            "target_words": target_words,
            "min_words": min_words,
            "max_words": max_words,
            "missing_h2": missing_h2,
            "invalid_external_citations": invalid_external_citations,
            "missing_personalization_signals": missing_personalization_signals,
        }

    def _ensure_required_headings(self, chapter: Dict[str, Any], content: str) -> str:
        missing_h2 = [
            sub for sub in chapter.get('subs', [])
            if not re.search(rf"(?im)^\s*##\s*{re.escape(sub)}\s*$", content)
        ]
        if not missing_h2:
            return content

        patched = content.rstrip()
        for heading in missing_h2:
            patched += (
                f"\n\n## {heading}\n"
                "### Rincian Inti\n"
                "1. Aktivitas utama pada bagian ini disesuaikan langsung dengan kebutuhan bisnis klien.\n"
                "- Risiko, dependensi, dan indikator keberhasilan dijabarkan secara terukur untuk eksekusi.\n"
            )
        return patched

    @staticmethod
    def _extract_json_object(text: str) -> Optional[Dict[str, Any]]:
        if not text:
            return None
        try:
            return json.loads(text)
        except Exception:
            pass
        match = re.search(r'\{.*\}', text, re.DOTALL)
        if not match:
            return None
        try:
            return json.loads(match.group(0))
        except Exception:
            return None

    @staticmethod
    def _chapter_excerpt(text: str, max_words: int = 170) -> str:
        words = re.findall(r"\S+", text)
        if len(words) <= max_words:
            return " ".join(words)
        head = " ".join(words[:110])
        tail = " ".join(words[-60:])
        return f"{head} ... {tail}"

    def _apply_global_coherence(
        self,
        chapter_outputs: Dict[str, str],
        selected_chapters: List[Dict[str, Any]],
        client: str,
        project: str
    ) -> Dict[str, str]:
        if len(chapter_outputs) < 2:
            return chapter_outputs

        snippets = []
        for chapter in selected_chapters:
            content = chapter_outputs.get(chapter['id'])
            if content:
                snippets.append(f"[{chapter['id']}] {chapter['title']} :: {self._chapter_excerpt(content)}")
        if len(snippets) < 2:
            return chapter_outputs

        snippets_text = "\n".join(snippets)
        prompt = f"""
        Audit koherensi proposal lintas bab untuk klien {client} dan inisiatif {project}.
        Ringkasan bab:
        {snippets_text}

        Keluarkan JSON murni tanpa markdown:
        {{
          "canonical_terms": [
            {{"preferred": "target state", "variants": ["kondisi target", "sasaran akhir"]}}
          ],
          "bridge_sentences": [
            {{"chapter_id": "c_2", "sentence": "Rumusan masalah ini menjadi dasar klasifikasi kebutuhan pada bab berikutnya."}}
          ]
        }}

        Aturan:
        - Maksimal 6 canonical_terms.
        - Maksimal 6 bridge_sentences.
        - sentence <= 25 kata.
        """
        try:
            res = self.ollama.chat(
                model=LLM_MODEL,
                messages=[
                    {'role': 'system', 'content': 'You output strictly valid JSON.'},
                    {'role': 'user', 'content': prompt}
                ],
                options={'num_ctx': 16384, 'num_predict': 900, 'temperature': 0.1}
            )
            directives = self._extract_json_object(res.get('message', {}).get('content', ''))
        except Exception:
            directives = None

        if not directives:
            return chapter_outputs

        canonical_terms = directives.get("canonical_terms", [])
        bridge_sentences = directives.get("bridge_sentences", [])
        revised = dict(chapter_outputs)

        for chap_id, text in revised.items():
            updated = text
            for item in canonical_terms:
                preferred = str(item.get("preferred", "")).strip()
                variants = item.get("variants", []) or []
                if not preferred:
                    continue
                for variant in variants:
                    variant_text = str(variant).strip()
                    if not variant_text or variant_text.lower() == preferred.lower():
                        continue
                    updated = re.sub(rf"(?i)\b{re.escape(variant_text)}\b", preferred, updated)
            revised[chap_id] = updated

        for item in bridge_sentences:
            chap_id = str(item.get("chapter_id", "")).strip()
            sentence = str(item.get("sentence", "")).strip()
            if not chap_id or chap_id not in revised or not sentence:
                continue
            if sentence in revised[chap_id]:
                continue
            revised[chap_id] = revised[chap_id].rstrip() + f"\n\n- {sentence}"

        return revised

    def _draft_chapter(
        self,
        chapter: Dict[str, Any],
        prompt: str,
        client: str,
        target_words: int,
        allowed_external_citations: Optional[Set[str]] = None,
        personalization_pack: Optional[Dict[str, Any]] = None
    ) -> str:
        allowed = set(allowed_external_citations or set())
        # First draft pass for a chapter.
        res = self.ollama.chat(
            model=LLM_MODEL,
            messages=[
                {'role': 'system', 'content': prompt},
                {'role': 'user', 'content': (
                    f"Tulis konten untuk {chapter['title']} dalam satu draft final. "
                    "Pastikan hard checks terpenuhi: H2 wajib lengkap, word range sesuai target, "
                    "ada numbered list dan bullet list, serta konten tetap konkret dan action-oriented."
                )}
            ],
            options={'num_ctx': 65536, 'num_predict': 4096, 'temperature': 0.25, 'top_p': 0.85, 'repeat_penalty': 1.1}
        )
        content = self._clean_external_citations((res.get('message', {}).get('content', '') or '').strip(), allowed)
        report = self._evaluate_chapter_quality(
            chapter,
            content,
            client,
            target_words=target_words,
            allowed_external_citations=allowed,
            personalization_pack=personalization_pack
        )
        hard_check_keys = {
            "missing_h2", "too_short", "too_long",
            "list_structure", "missing_visual", "citation_policy", "missing_personalization"
        }
        hard_issues = [i for i in report["issues"] if i in hard_check_keys]
        if not hard_issues:
            return content

        # Retry once when hard checks fail.
        citation_policy_note = (
            f"Allowed external citations: {', '.join(sorted(allowed)) if allowed else 'none'}.\n"
            f"Invalid external citations found: {', '.join(report.get('invalid_external_citations', [])) or '-'}.\n"
            "Hapus semua sitasi eksternal yang tidak ada di daftar allowed.\n"
        )
        personalization_note = (
            f"Sinyal personalisasi yang belum muncul: "
            f"{', '.join(report.get('missing_personalization_signals', [])) or '-'}.\n"
            "Pastikan konten mencerminkan konteks klien, KPI tailoring, terminologi domain, dan anchor inisiatif.\n"
        )
        retry_prompt = (
            f"Perbaiki draft {chapter['title']} agar lulus hard quality checks.\n"
            f"Issues: {', '.join(hard_issues)}\n"
            f"Word count: {report['word_count']} (target {report['target_words']}, range {report['min_words']}-{report['max_words']}).\n"
            f"Missing H2: {', '.join(report['missing_h2']) if report['missing_h2'] else '-'}\n"
            f"{citation_policy_note}"
            f"{personalization_note}"
            "Pertahankan fakta dan konteks. Keluarkan versi final saja.\n\n"
            f"DRAFT:\n{content}"
        )
        try:
            retry = self.ollama.chat(
                model=LLM_MODEL,
                messages=[
                    {'role': 'system', 'content': prompt},
                    {'role': 'user', 'content': retry_prompt}
                ],
                options={'num_ctx': 65536, 'num_predict': 4096, 'temperature': 0.2, 'top_p': 0.85, 'repeat_penalty': 1.1}
            )
            improved = (retry.get('message', {}).get('content', '') or '').strip()
            if improved:
                content = self._clean_external_citations(improved, allowed)
        except Exception:
            pass

        final_report = self._evaluate_chapter_quality(
            chapter,
            content,
            client,
            target_words=target_words,
            allowed_external_citations=allowed,
            personalization_pack=personalization_pack
        )
        if final_report.get("missing_h2"):
            content = self._ensure_required_headings(chapter, content)
            final_report = self._evaluate_chapter_quality(
                chapter,
                content,
                client,
                target_words=target_words,
                allowed_external_citations=allowed,
                personalization_pack=personalization_pack
            )
        if "citation_policy" in final_report.get("issues", []):
            content = self._clean_external_citations(content, allowed)
            final_report = self._evaluate_chapter_quality(
                chapter,
                content,
                client,
                target_words=target_words,
                allowed_external_citations=allowed,
                personalization_pack=personalization_pack
            )

        if final_report["issues"]:
            logger.warning(f"Quality checks not fully satisfied for {chapter['title']}: {', '.join(final_report['issues'])}")
        return content

    def _tighten_chapter(
        self,
        chapter: Dict[str, Any],
        prompt: str,
        content: str,
        target_words: int,
        allowed_external_citations: Optional[Set[str]] = None
    ) -> str:
        allowed = set(allowed_external_citations or set())
        try:
            res = self.ollama.chat(
                model=LLM_MODEL,
                messages=[
                    {'role': 'system', 'content': prompt},
                    {'role': 'user', 'content': (
                        f"Rapikan dan padatkan konten {chapter['title']} menjadi sekitar {target_words} kata. "
                        "Pertahankan semua heading H2 wajib, poin kunci, dan keterbacaan eksekutif. "
                        "Hapus repetisi dan kalimat pengisi.\n\n"
                        f"KONTEN SAAT INI:\n{content}"
                    )}
                ],
                options={'num_ctx': 65536, 'num_predict': 4096, 'temperature': 0.15, 'top_p': 0.8, 'repeat_penalty': 1.1}
            )
            revised = (res.get('message', {}).get('content', '') or '').strip()
            return self._clean_external_citations(revised or content, allowed)
        except Exception:
            return self._clean_external_citations(content, allowed)

    def _fit_into_word_budget(
        self,
        chapter_outputs: Dict[str, str],
        chapter_prompts: Dict[str, str],
        chapter_map: Dict[str, Dict[str, Any]],
        chapter_targets: Dict[str, int],
        max_words: int,
        allowed_external_citations: Optional[Set[str]] = None
    ) -> Dict[str, str]:
        outputs = dict(chapter_outputs)

        def total_words() -> int:
            return sum(self._word_count(text) for text in outputs.values() if text)

        current = total_words()
        if current <= max_words:
            return outputs

        for _ in range(3):
            current = total_words()
            if current <= max_words:
                break

            overflow = current - max_words
            changed = False
            chapter_order = sorted(
                [cid for cid, text in outputs.items() if text],
                key=lambda cid: self._word_count(outputs[cid]),
                reverse=True
            )

            for cid in chapter_order:
                if overflow <= 0:
                    break
                chapter = chapter_map.get(cid)
                if not chapter:
                    continue
                prompt = chapter_prompts.get(cid)
                if not prompt:
                    continue

                current_words = self._word_count(outputs[cid])
                minimum = self._chapter_floor_words(cid, for_compression=True)
                reducible = max(0, current_words - minimum)
                if reducible < 80:
                    continue

                target = max(minimum, current_words - min(reducible, max(120, overflow)))
                target = min(target, max(minimum, int(chapter_targets.get(cid, target) * 0.95)))
                compressed = self._tighten_chapter(
                    chapter,
                    prompt,
                    outputs[cid],
                    target,
                    allowed_external_citations=allowed_external_citations
                )
                if compressed == outputs[cid]:
                    continue
                outputs[cid] = compressed
                overflow = total_words() - max_words
                changed = True

            if not changed:
                break

        if total_words() > max_words:
            ratio = max_words / max(total_words(), 1)
            for cid, text in list(outputs.items()):
                chapter = chapter_map.get(cid)
                prompt = chapter_prompts.get(cid)
                if not chapter or not prompt or not text:
                    continue
                minimum = self._chapter_floor_words(cid, for_compression=True)
                current_words = self._word_count(text)
                target = max(minimum, int(current_words * ratio))
                if target >= current_words - 60:
                    continue
                outputs[cid] = self._tighten_chapter(
                    chapter,
                    prompt,
                    text,
                    target,
                    allowed_external_citations=allowed_external_citations
                )

        return outputs

    def generate_document(
        self,
        client: str,
        project: str,
        budget: str,
        service_type: str,
        project_goal: str,
        project_type: str,
        timeline: str,
        notes: str,
        regulations: str,
        chapter_id: Optional[str] = None
    ) -> Tuple[Document, str]:
        selected_chapters = self._resolve_chapters(chapter_id)
        chapter_targets = self._chapter_word_targets(selected_chapters)
        content_word_budget = self._content_word_budget()

        use_demo_logic = self.firm_api.uses_demo_logic()
        firm_data = self.firm_api.get_project_standards(project_type)
        firm_profile = self.firm_api.get_firm_profile()
        base_client = re.sub(r'\b(Cabang|Branch|Tbk)\b.*$|^(PT\.|CV\.)', '', client, flags=re.IGNORECASE).strip()
        relationship_context = self.firm_api.get_client_relationship(base_client)
        research_bundle = self._get_research_bundle(
            base_client,
            regulations,
            include_collaboration=use_demo_logic
        )
        if not use_demo_logic:
            research_bundle = dict(research_bundle)
            research_bundle["collaboration"] = relationship_context.get("summary", "")
        personalization_pack = self._build_personalization_pack(
            client=client,
            project=project,
            project_goal=project_goal,
            timeline=timeline,
            notes=notes,
            regulations=regulations,
            research_bundle=research_bundle,
            relationship_context=relationship_context
        )
        allowed_external_citations = self._collect_allowed_external_citations(research_bundle)
        proposal_contract = self._build_proposal_contract(
            client=client,
            project=project,
            budget=budget,
            service_type=service_type,
            project_goal=project_goal,
            project_type=project_type,
            timeline=timeline,
            notes=notes,
            regulations=regulations,
            selected_chapters=selected_chapters,
            research_bundle=research_bundle,
            firm_data=firm_data,
            personalization_pack=personalization_pack
        )
        logo_future = self.io_pool.submit(LogoManager.get_logo_and_color, base_client)

        context_futures = {
            chapter['id']: self.io_pool.submit(
                self._build_chapter_prompt,
                chapter, client, project, budget, service_type, project_goal, project_type, timeline,
                notes, regulations, firm_data, firm_profile, research_bundle, personalization_pack, proposal_contract,
                chapter_targets.get(chapter['id'], self._target_words(chapter))
            )
            for chapter in selected_chapters
        }

        chapter_map = {chapter['id']: chapter for chapter in selected_chapters}
        chapter_prompts: Dict[str, str] = {}
        chapter_outputs: Dict[str, str] = {}
        for chapter in selected_chapters:
            ctx = context_futures[chapter['id']].result()
            if not ctx['success']:
                continue
            chapter_prompts[chapter['id']] = ctx['prompt']
            try:
                chapter_outputs[chapter['id']] = self._draft_chapter(
                    chapter=chapter,
                    prompt=ctx['prompt'],
                    client=client,
                    target_words=chapter_targets.get(chapter['id'], self._target_words(chapter)),
                    allowed_external_citations=allowed_external_citations,
                    personalization_pack=personalization_pack
                )
            except Exception as e:
                logger.error(f"Generation Error for {chapter['title']}: {e}")

        chapter_outputs = {
            chapter_id: self._clean_external_citations(content, allowed_external_citations)
            for chapter_id, content in chapter_outputs.items()
        }

        chapter_outputs = self._fit_into_word_budget(
            chapter_outputs=chapter_outputs,
            chapter_prompts=chapter_prompts,
            chapter_map=chapter_map,
            chapter_targets=chapter_targets,
            max_words=content_word_budget,
            allowed_external_citations=allowed_external_citations
        )
        chapter_outputs = self._apply_global_coherence(chapter_outputs, selected_chapters, client, project)
        chapter_outputs = {
            chapter_id: self._clean_external_citations(content, allowed_external_citations)
            for chapter_id, content in chapter_outputs.items()
        }
        chapter_outputs = self._fit_into_word_budget(
            chapter_outputs=chapter_outputs,
            chapter_prompts=chapter_prompts,
            chapter_map=chapter_map,
            chapter_targets=chapter_targets,
            max_words=content_word_budget,
            allowed_external_citations=allowed_external_citations
        )
        generated_words = sum(self._word_count(t) for t in chapter_outputs.values() if t)
        estimated_pages = self._estimated_pages(generated_words)
        if estimated_pages > MAX_PROPOSAL_PAGES:
            logger.warning(
                f"Estimated page count is above limit ({estimated_pages}>{MAX_PROPOSAL_PAGES}); applying one more compacting pass."
            )
            tighter_budget = max(int(content_word_budget * 0.94), 4000)
            chapter_outputs = self._fit_into_word_budget(
                chapter_outputs=chapter_outputs,
                chapter_prompts=chapter_prompts,
                chapter_map=chapter_map,
                chapter_targets=chapter_targets,
                max_words=tighter_budget,
                allowed_external_citations=allowed_external_citations
            )

        if chapter_outputs.get("c_closing"):
            chapter_outputs["c_closing"] = self._inject_verified_firm_contact(
                chapter_outputs["c_closing"],
                firm_profile
            )

        try:
            logo_stream, theme_color = logo_future.result(timeout=8)
        except Exception:
            logo_stream, theme_color = None, DEFAULT_COLOR

        doc = Document()
        StyleEngine.apply_document_styles(doc)
        
        # Cover page.
        for _ in range(2):
            doc.add_paragraph()

        if logo_stream:
            try:
                logo_stream.seek(0)
                cover_logo = doc.add_paragraph()
                cover_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cover_logo.add_run().add_picture(logo_stream, width=Inches(1.8))
            except (UnrecognizedImageError, OSError, ValueError) as e:
                logger.warning(f"Logo skipped due to unsupported image format: {e}")

        title = doc.add_paragraph("ARSITEKTUR PROPOSAL")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.bold = True
        title_run.font.size = Pt(30)
        title_run.font.color.rgb = RGBColor(*theme_color)

        if len(selected_chapters) == 1:
            subtitle_text = f"{selected_chapters[0]['title']} ({service_type} – {project_type})"
        else:
            subtitle_text = f"{service_type} – {project_type}"

        subtitle = doc.add_paragraph(subtitle_text)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].font.size = Pt(14)

        doc.add_paragraph()
        client_line = doc.add_paragraph(f"Untuk: {client}")
        client_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        client_line.runs[0].bold = True
        client_line.runs[0].font.size = Pt(16)

        project_line = doc.add_paragraph(f"Inisiatif: {project}")
        project_line.alignment = WD_ALIGN_PARAGRAPH.CENTER

        meta = doc.add_paragraph(
            f"Durasi: {timeline} | Estimasi Investasi: {budget or 'Menyesuaikan ruang lingkup'} | Tanggal: {datetime.now().strftime('%d %B %Y')}"
        )
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

        firm_contact_lines = FirmAPIClient.build_contact_lines(firm_profile)
        contact_text = f"Disusun oleh {WRITER_FIRM_NAME}"
        if firm_contact_lines:
            contact_text += "\n" + "\n".join(firm_contact_lines)
        contact = doc.add_paragraph(contact_text)
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_page_break()

        rendered_any = False
        for i, chapter in enumerate(selected_chapters):
            content = chapter_outputs.get(chapter['id'], '').strip()
            if not content:
                continue
            rendered_any = True
            h = doc.add_heading(chapter['title'], level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            h.runs[0].font.color.rgb = RGBColor(*theme_color)
            DocumentBuilder.process_content(doc, content, theme_color, chapter['title'])

            has_next = any(
                chapter_outputs.get(next_chapter['id'], '').strip()
                for next_chapter in selected_chapters[i + 1:]
            )
            if has_next:
                doc.add_page_break()

        if not rendered_any:
            doc.add_paragraph("Konten proposal belum berhasil digenerate. Mohon ulangi proses.")

        base_name = f"Proposal_{client}_{project}"
        if len(selected_chapters) == 1:
            chapter_slug = re.sub(r'[^A-Za-z0-9]+', '_', selected_chapters[0]['title']).strip('_')
            base_name = f"{base_name}_{chapter_slug}"

        return doc, base_name.replace(" ", "_")


    def run(
        self,
        client: str,
        project: str,
        budget: str,
        service_type: str,
        project_goal: str,
        project_type: str,
        timeline: str,
        notes: str,
        regulations: str,
        chapter_id: Optional[str] = None
    ) -> Tuple[Document, str]:
        return self.generate_document(
            client=client,
            project=project,
            budget=budget,
            service_type=service_type,
            project_goal=project_goal,
            project_type=project_type,
            timeline=timeline,
            notes=notes,
            regulations=regulations,
            chapter_id=chapter_id,
        )
